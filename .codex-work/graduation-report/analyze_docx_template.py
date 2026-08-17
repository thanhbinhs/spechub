from __future__ import annotations

import hashlib
import json
import sys
import zipfile
from collections import Counter
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def emu_to_inches(value):
    return None if value is None else round(value / 914400, 4)


def points(value):
    return None if value is None else round(value.pt, 2)


def rgb(value):
    return None if value is None else str(value)


def style_info(style):
    pf = style.paragraph_format
    font = style.font
    rfonts = style.element.xpath(".//w:rPr/w:rFonts")
    return {
        "name": style.name,
        "type": str(style.type),
        "base_style": style.base_style.name if style.base_style else None,
        "font": {
            "name": font.name,
            "ascii": rfonts[0].get(qn("w:ascii")) if rfonts else None,
            "east_asia": rfonts[0].get(qn("w:eastAsia")) if rfonts else None,
            "size_pt": points(font.size),
            "bold": font.bold,
            "italic": font.italic,
            "color": rgb(font.color.rgb),
        },
        "paragraph": {
            "alignment": str(pf.alignment),
            "space_before_pt": points(pf.space_before),
            "space_after_pt": points(pf.space_after),
            "line_spacing": str(pf.line_spacing),
            "line_spacing_rule": str(pf.line_spacing_rule),
            "left_indent_in": emu_to_inches(pf.left_indent),
            "right_indent_in": emu_to_inches(pf.right_indent),
            "first_line_indent_in": emu_to_inches(pf.first_line_indent),
            "keep_with_next": pf.keep_with_next,
            "keep_together": pf.keep_together,
            "page_break_before": pf.page_break_before,
        },
    }


def paragraph_sample(paragraph):
    return {
        "style": paragraph.style.name if paragraph.style else None,
        "text": paragraph.text[:240],
        "alignment": str(paragraph.alignment),
        "space_before_pt": points(paragraph.paragraph_format.space_before),
        "space_after_pt": points(paragraph.paragraph_format.space_after),
        "line_spacing": str(paragraph.paragraph_format.line_spacing),
        "first_line_indent_in": emu_to_inches(
            paragraph.paragraph_format.first_line_indent
        ),
        "left_indent_in": emu_to_inches(paragraph.paragraph_format.left_indent),
        "runs": [
            {
                "text": run.text[:120],
                "font": run.font.name,
                "size_pt": points(run.font.size),
                "bold": run.bold,
                "italic": run.italic,
                "color": rgb(run.font.color.rgb),
            }
            for run in paragraph.runs[:6]
        ],
    }


def analyze(path: Path, render_dir: Path):
    doc = Document(path)
    page_count = len(list(render_dir.glob("page-*.png")))
    style_counts = Counter(
        p.style.name if p.style else "(none)" for p in doc.paragraphs
    )
    non_empty = [p for p in doc.paragraphs if p.text.strip()]
    headings = [
        {
            "index": index,
            "style": p.style.name if p.style else None,
            "text": p.text.strip(),
        }
        for index, p in enumerate(doc.paragraphs)
        if p.style and p.style.name.startswith("Heading")
    ]
    samples_by_style = {}
    for p in non_empty:
        name = p.style.name if p.style else "(none)"
        samples_by_style.setdefault(name, [])
        if len(samples_by_style[name]) < 4:
            samples_by_style[name].append(paragraph_sample(p))

    desired_styles = [
        "Normal",
        "Title",
        "Subtitle",
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "Heading 4",
        "Caption",
        "TOC 1",
        "TOC 2",
        "TOC 3",
        "List Paragraph",
    ]
    styles = {
        name: style_info(doc.styles[name])
        for name in desired_styles
        if name in doc.styles
    }
    sections = []
    for index, section in enumerate(doc.sections, start=1):
        sections.append(
            {
                "index": index,
                "start_type": str(section.start_type),
                "orientation": str(section.orientation),
                "page_width_in": emu_to_inches(section.page_width),
                "page_height_in": emu_to_inches(section.page_height),
                "top_margin_in": emu_to_inches(section.top_margin),
                "bottom_margin_in": emu_to_inches(section.bottom_margin),
                "left_margin_in": emu_to_inches(section.left_margin),
                "right_margin_in": emu_to_inches(section.right_margin),
                "header_distance_in": emu_to_inches(section.header_distance),
                "footer_distance_in": emu_to_inches(section.footer_distance),
                "different_first_page_header_footer": section.different_first_page_header_footer,
            }
        )

    package_parts = []
    with zipfile.ZipFile(path) as archive:
        for info in sorted(archive.infolist(), key=lambda value: value.filename):
            payload = archive.read(info.filename)
            package_parts.append(
                {
                    "path": info.filename,
                    "size": info.file_size,
                    "sha256": hashlib.sha256(payload).hexdigest(),
                }
            )

    return {
        "reference": str(path.resolve()),
        "sha256": hashlib.sha256(path.read_bytes()).hexdigest(),
        "size_bytes": path.stat().st_size,
        "page_count": page_count,
        "section_count": len(doc.sections),
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "inline_shape_count": len(doc.inline_shapes),
        "section_geometry": sections,
        "style_counts": dict(style_counts.most_common()),
        "styles": styles,
        "headings": headings,
        "samples_by_style": samples_by_style,
        "package_parts": package_parts,
    }


def main():
    if len(sys.argv) != 4:
        raise SystemExit("usage: analyze_docx_template.py DOCX RENDER_DIR OUT_JSON")
    path = Path(sys.argv[1])
    render_dir = Path(sys.argv[2])
    out = Path(sys.argv[3])
    out.write_text(
        json.dumps(analyze(path, render_dir), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


if __name__ == "__main__":
    main()
