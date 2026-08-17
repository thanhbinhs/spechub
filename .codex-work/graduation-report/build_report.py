from __future__ import annotations

import argparse
import json
import posixpath
import re
from datetime import datetime
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageOps
from lxml import etree


ROOT = Path(__file__).resolve().parent
WORKSPACE = ROOT.parents[1]
ASSETS = ROOT / "assets"
DIAGRAMS = ASSETS / "diagrams"
UI = ASSETS / "ui-bw"
OUTPUT = ROOT / "output"
OUTPUT.mkdir(parents=True, exist_ok=True)

TEMPLATE = ROOT / "templates" / "21012899_PhamThanhTrung.docx"
EVIDENCE_PATH = ROOT / "analysis" / "spechub-evidence.json"
SCHEMA_PATH = WORKSPACE / "packages" / "database" / "prisma" / "schema.prisma"

NAVY = "000000"
BLUE = "000000"
LIGHT_BLUE = "E6E6E6"
HEADER_FILL = "D9D9D9"
ALT_FILL = "F2F2F2"
GRAY = "595959"
WHITE = "FFFFFF"
BLACK = "000000"

PAGE_MAP: dict[str, str] = {}
MARKERS = False
FIGURE_ENTRIES: list[tuple[str, str]] = []
TABLE_ENTRIES: list[tuple[str, str]] = []
HEADING_ENTRIES: list[tuple[int, str, str]] = []


ACADEMIC_REPLACEMENTS = [
    ("Prisma schema", "lược đồ Prisma"), ("schema Prisma", "lược đồ Prisma"),
    ("Schema", "Lược đồ"),
    ("route filesystem", "tuyến theo cấu trúc thư mục"),
    ("ERD nhóm", "Sơ đồ thực thể–liên kết nhóm"),
    ("ERD tiêu biểu", "Sơ đồ thực thể–liên kết tiêu biểu"),
    ("Use Case", "ca sử dụng"), ("Use case", "ca sử dụng"), ("use case", "ca sử dụng"),
    ("Controller inventory", "Kết quả thống kê các bộ điều khiển"),
    ("controller inventory", "thống kê các bộ điều khiển"),
    ("Static inventory", "Thống kê mã nguồn"), ("inventory", "thống kê"),
    ("Bằng chứng kiểm thử và kiểm chứng", "Kết quả kiểm thử và đánh giá"),
    ("Bằng chứng hiện thực", "Căn cứ đánh giá"),
    ("bằng chứng hiện thực", "căn cứ đánh giá"),
    ("bằng chứng thực tế", "kết quả quan sát thực tế"),
    ("snapshot", "ảnh chụp"), ("Snapshot", "Ảnh chụp"), ("mockup", "mẫu mô phỏng"),
    ("workspace root", "thư mục gốc của dự án"), ("workspace", "không gian làm việc"),
    ("web route", "tuyến giao diện"), ("route handler", "hàm xử lý tuyến"),
    ("Controller", "Bộ điều khiển"), ("controller", "bộ điều khiển"),
    ("controllers", "các bộ điều khiển"),
    ("Endpoint", "Điểm cuối API"), ("endpoint", "điểm cuối API"),
    ("Route", "Tuyến"), ("route", "tuyến"),
    ("pipeline", "quy trình xử lý"), ("Pipeline", "Quy trình xử lý"),
    ("Worker", "Tiến trình nền"), ("worker", "tiến trình nền"),
    ("service", "dịch vụ"), ("Service", "Dịch vụ"),
    ("services", "các dịch vụ"),
    ("module", "mô-đun"), ("Module", "Mô-đun"),
    ("package", "gói"), ("Package", "Gói"),
    ("schema validation", "kiểm tra lược đồ"), ("schema", "lược đồ"),
    ("validate", "kiểm tra tính hợp lệ"), ("Validation", "Kiểm tra tính hợp lệ"),
    ("validation", "kiểm tra tính hợp lệ"), ("authorize", "kiểm tra quyền"),
    ("citation coverage", "độ phủ trích dẫn"),
    ("Citation", "Trích dẫn nguồn"), ("citation", "trích dẫn nguồn"),
    ("published", "đã xuất bản"), ("draft", "bản nháp"),
    ("raw page", "trang dữ liệu thô"), ("Raw page", "Trang dữ liệu thô"),
    ("Data source", "Nguồn dữ liệu"), ("Parser", "Bộ phân tích dữ liệu"),
    ("human review", "đánh giá của chuyên gia"),
    ("threat-model review", "rà soát mô hình đe dọa"),
    ("dependency scanning", "quét lỗ hổng thư viện phụ thuộc"),
    ("secret scanning", "quét thông tin bí mật"),
    ("backup/restore drill", "diễn tập sao lưu và khôi phục"),
    ("event contract", "hợp đồng sự kiện"),
    ("modular monolith", "kiến trúc nguyên khối mô-đun"),
    ("review queue", "hàng đợi kiểm duyệt"), ("review state", "trạng thái kiểm duyệt"),
    ("review", "kiểm duyệt"), ("Review", "Kiểm duyệt"), ("diff", "sai khác"),
    ("Audit", "Nhật ký hoạt động"), ("audit", "nhật ký hoạt động"),
    ("API key", "khóa API"), ("rate limit", "giới hạn tần suất"),
    ("Usage", "Mức sử dụng"),
    ("scope", "phạm vi quyền"), ("quota", "hạn mức"), ("usage", "mức sử dụng"),
    ("Wishlist", "Danh sách yêu thích"), ("wishlist", "danh sách yêu thích"),
    ("Notification", "Thông báo"), ("notification", "thông báo"),
    ("delivery", "lần gửi"), ("Alert", "Cảnh báo"), ("alert", "cảnh báo"),
    ("refresh session", "phiên làm mới"), ("refresh token", "mã làm mới"),
    ("access token", "mã truy cập"), ("session", "phiên"),
    ("structured log", "nhật ký có cấu trúc"), ("request ID", "mã định danh yêu cầu"),
    ("Request ID", "Mã định danh yêu cầu"),
    ("latency", "độ trễ"), ("readiness", "mức sẵn sàng"), ("liveness", "trạng thái tiến trình"),
    ("cache", "bộ nhớ đệm"), ("fallback", "phương án dự phòng"),
    ("Cache", "Bộ nhớ đệm"), ("Readiness", "Kiểm tra mức sẵn sàng"),
    ("retry", "thử lại"), ("stale", "quá hạn"),
    ("idempotency", "tính lũy đẳng"), ("idempotent", "lũy đẳng"),
    ("parser", "bộ phân tích dữ liệu"), ("fetch", "thu thập"),
    ("candidate", "dữ liệu đề xuất"), ("content hash", "mã băm nội dung"),
    ("top-k context", "k ngữ cảnh phù hợp nhất"), ("context", "ngữ cảnh"),
    ("retrieval", "truy xuất"), ("generation", "sinh nội dung"),
    ("presentation", "trình bày"), ("prompt injection", "tấn công chèn chỉ dẫn"),
    ("prompt", "chỉ dẫn"), ("embedding", "véc-tơ nhúng"),
    ("unit test", "kiểm thử đơn vị"), ("test unit", "kiểm thử đơn vị"),
    ("UI smoke", "Kiểm tra nhanh giao diện"),
    ("Playwright E2E", "kiểm thử đầu cuối bằng Playwright"),
    ("E2E assertion", "khẳng định kiểm thử đầu cuối"),
    ("E2E toàn tuyến", "kiểm thử đầu cuối toàn tuyến"),
    ("kiểm thử E2E", "kiểm thử đầu cuối (E2E)"),
    ("pentest", "kiểm thử xâm nhập"), ("benchmark", "đánh giá chuẩn"),
    ("baseline", "mốc tham chiếu"), ("commit/worktree", "phiên bản mã nguồn"),
    ("controller/service", "bộ điều khiển/dịch vụ"),
    ("current revision", "phiên bản hiện hành"),
    ("current_revision", "phiên bản hiện hành"),
    ("request HTTP", "yêu cầu HTTP"),
    ("public endpoint", "điểm cuối API công khai"),
    ("public decorator", "khai báo công khai"),
    ("public", "công khai"),
    ("backlog", "danh sách công việc"), ("freshness", "độ cập nhật"),
    ("conflict rate", "tỷ lệ xung đột"),
    ("golden set", "bộ dữ liệu đánh giá chuẩn"), ("feedback", "phản hồi"),
    ("variant", "biến thể"), ("alias", "tên thay thế"),
    ("target price", "mức giá mục tiêu"), ("quiet hours", "khung giờ không làm phiền"),
    ("production", "môi trường vận hành"),
    ("request", "yêu cầu"), ("Search", "Tìm kiếm"), ("search", "tìm kiếm"),
    ("Role", "Vai trò"), ("role", "vai trò"),
    ("category", "loại thiết bị"),
    ("workflow", "quy trình nghiệp vụ"),
    ("subscription", "gói thuê bao"), ("billing", "thanh toán"),
    ("affiliate", "tiếp thị liên kết"), ("moderation", "kiểm duyệt"),
    ("health/metrics", "kiểm tra trạng thái và chỉ số giám sát"),
    ("Health", "Trạng thái hoạt động"), ("health", "trạng thái hoạt động"),
    ("metrics", "chỉ số giám sát"),
    ("server-side", "phía máy chủ"), ("server", "máy chủ"),
    ("Server rendering", "Kết xuất phía máy chủ"),
    ("Client component", "Thành phần phía máy khách"),
    ("Client", "Máy khách"), ("client", "máy khách"),
    ("Input", "Dữ liệu đầu vào"), ("input", "dữ liệu đầu vào"),
    ("metadata", "siêu dữ liệu"), ("hydrate", "truy xuất chi tiết"),
    ("Log", "Nhật ký"), ("log", "nhật ký"),
    ("namespace", "không gian tên"), ("dead-letter", "hàng đợi lỗi"),
    ("ownership", "quyền sở hữu"), ("rotate", "luân chuyển"),
    ("scrub", "loại bỏ dữ liệu nhạy cảm khỏi"),
    ("model adapter", "bộ điều hợp mô hình"), ("Model", "Mô hình"), ("model", "mô hình"),
    ("marketing", "tiếp thị"), ("media", "dữ liệu đa phương tiện"),
    ("status code", "mã trạng thái"), ("status", "trạng thái"),
    ("currency", "tiền tệ"), ("pagination", "phân trang"),
    ("timeout", "thời gian chờ"),
    ("data dictionary", "từ điển dữ liệu"),
    ("Session", "Phiên"), ("Alias", "Tên thay thế"),
    ("Runtime", "Môi trường thực thi"),
    ("Ownership", "Quyền sở hữu"), ("Subscription", "Gói thuê bao"),
    ("Billing", "Thanh toán"), ("Idempotency", "Tính lũy đẳng"),
    ("Latency", "Độ trễ"),
    ("Moderator", "Kiểm duyệt viên"), ("Editor", "Biên tập viên"),
    ("Admin", "Quản trị viên"), ("Revision", "Phiên bản"),
    ("revision", "phiên bản"), ("Transaction", "Giao dịch"),
    ("transaction", "giao dịch"), ("Gateway", "Cổng API"),
    ("guard", "bộ bảo vệ"), ("Key", "Khóa"), ("key", "khóa"),
    ("header", "phần đầu phản hồi"), ("versioned", "được quản lý phiên bản"),
    ("guidance", "hướng dẫn"), ("versioning", "quản lý phiên bản"),
    ("job", "tác vụ"), ("reindex", "lập chỉ mục lại"),
    ("rollback", "hoàn tác"), ("hash", "mã băm"),
    ("decimal", "số thập phân"), ("preference", "tùy chọn"),
    ("unsubscribe", "hủy đăng ký"), ("Unique", "Ràng buộc duy nhất"),
    ("Optimistic concurrency", "Kiểm soát đồng thời lạc quan"),
    ("sanitize rich text", "làm sạch văn bản định dạng"),
    ("merge", "hợp nhất"), ("actor", "tác nhân"),
]


# Giữ nguyên các technical terms phổ biến theo yêu cầu của báo cáo. Danh sách
# này thay thế cơ chế Việt hóa cũ; phần diễn giải vẫn viết bằng tiếng Việt,
# nhưng tên artifact, component, protocol, test level và software concept được
# trình bày bằng thuật ngữ gốc để tránh sai nghĩa chuyên môn.
TECHNICAL_TERM_NORMALIZATION = [
    ("Sơ đồ thực thể–liên kết", "ERD"),
    ("Bộ nhớ đệm", "Cache"),
    ("Tiến trình nền", "Worker"),
    ("Mức sẵn sàng", "Readiness"),
    ("Bộ điều khiển", "Controller"),
    ("Điểm cuối API", "API endpoint"),
    ("Kiểm thử đầu cuối toàn tuyến", "End-to-end test (E2E)"),
    ("Kiểm thử đầu cuối", "End-to-end test (E2E)"),
    ("Kiểm thử đơn vị", "Unit test"),
    ("Kiểm tra quy tắc mã", "Lint"),
    ("Kiểm tra kiểu dữ liệu", "Type-check"),
    ("Lược đồ hợp lệ", "Schema hợp lệ"),
    ("kiểm thử xâm nhập", "penetration test"),
    ("kiểm thử tải", "load test"),
    ("kiểm thử giao diện đầu cuối", "end-to-end UI test (E2E)"),
    ("khả năng biên dịch", "build"),
    ("biên dịch toàn dự án", "project build"),
    ("Biên dịch toàn dự án", "Project build"),
    ("kiểm thử", "test"),
    ("Kiểm thử", "Test"),
    ("biên dịch", "build"),
    ("Biên dịch", "Build"),
    ("mô-đun", "module"),
    ("Mô-đun", "Module"),
    ("tác vụ nền", "background job"),
    ("phiên đăng nhập", "session"),
    ("hàng đợi", "queue"),
    ("Hàng đợi", "Queue"),
    ("độ phủ trích dẫn", "citation coverage"),
    ("trích dẫn nguồn", "citation"),
    ("trích dẫn", "citation"),
    ("Trích dẫn", "Citation"),
    ("truy xuất tăng cường", "Retrieval-Augmented Generation (RAG)"),
    ("Truy xuất tăng cường", "Retrieval-Augmented Generation (RAG)"),
    ("truy xuất", "retrieval"),
    ("Truy xuất", "Retrieval"),
    ("sinh nội dung", "generation"),
    ("lập chỉ mục", "indexing"),
    ("Lập chỉ mục", "Indexing"),
    ("tìm kiếm ngữ nghĩa", "semantic search"),
    ("ngữ cảnh", "context"),
    ("mật khẩu băm", "password hash"),
    ("mật khẩu được băm", "password được lưu dưới dạng hash"),
    ("phân trang", "pagination"),
    ("khóa ngoại", "foreign key"),
    ("Khóa ngoại", "Foreign key"),
    ("khóa chính", "primary key"),
    ("giao dịch cơ sở dữ liệu", "database transaction"),
    ("hợp đồng dữ liệu", "data contract"),
    ("bảng nối", "join table"),
    ("Bảng nối", "Join table"),
    ("lược đồ", "schema"),
    ("Lược đồ", "Schema"),
    ("mô hình Prisma", "Prisma model"),
    ("Mô hình Prisma", "Prisma model"),
    ("mô hình dữ liệu", "data model"),
    ("chỉ mục", "index"),
    ("Chỉ mục", "Index"),
    ("cơ sở dữ liệu", "database"),
    ("Cơ sở dữ liệu", "Database"),
    ("thời gian chờ", "timeout"),
    ("độ cập nhật", "freshness"),
    ("tỷ lệ xung đột", "conflict rate"),
    ("từ điển dữ liệu", "data dictionary"),
    ("mức độ ưu tiên", "priority"),
    ("kiểm tra quyền", "authorization"),
    ("xác thực", "authentication"),
    ("Xác thực", "Authentication"),
    ("phân quyền", "authorization"),
    ("Phân quyền", "Authorization"),
    ("phía máy chủ", "server-side"),
    ("phía máy khách", "client-side"),
    ("trình khách Prisma", "Prisma Client"),
    ("phần đầu phản hồi", "response header"),
    ("ngoại tuyến", "offline"),
    ("bảo mật", "security"),
    ("Bảo mật", "Security"),
    ("quan sát hệ thống", "observability"),
    ("Quan sát hệ thống", "Observability"),
    ("khả năng phục hồi", "resilience"),
    ("lớp trình bày", "presentation layer"),
    ("Lớp trình bày", "Presentation layer"),
    ("lớp truy cập dữ liệu", "data access layer"),
    ("dữ liệu quan hệ", "relational data"),
    ("Sơ đồ thực thể–liên kết", "ERD"),
    ("sơ đồ thực thể–liên kết", "ERD"),
    ("kiến trúc nguyên khối mô-đun", "modular monolith"),
    ("kiểm thử đầu cuối toàn tuyến", "end-to-end test (E2E)"),
    ("kiểm thử đầu cuối bằng Playwright", "Playwright end-to-end test (E2E)"),
    ("kiểm thử đầu cuối", "end-to-end test (E2E)"),
    ("kiểm thử đơn vị", "unit test"),
    ("kiểm tra quy tắc viết mã", "lint"),
    ("kiểm tra quy tắc mã", "lint"),
    ("kiểm tra kiểu dữ liệu", "type-check"),
    ("kiểm tra lược đồ", "schema validation"),
    ("kiểm tra mức sẵn sàng", "readiness check"),
    ("mức sẵn sàng", "readiness"),
    ("trạng thái tiến trình", "liveness"),
    ("chỉ số giám sát", "metrics"),
    ("bộ nhớ đệm", "cache"),
    ("công cụ tìm kiếm", "search engine"),
    ("hàng đợi kiểm duyệt", "review queue"),
    ("hàng đợi lỗi", "dead-letter queue"),
    ("nhật ký hoạt động", "audit log"),
    ("nhật ký kiểm toán", "audit log"),
    ("tiến trình nền", "worker"),
    ("gói dùng chung", "shared package"),
    ("kho mã nguồn", "repository"),
    ("thư mục gốc của dự án", "workspace root"),
    ("không gian làm việc", "workspace"),
    ("tuyến theo cấu trúc thư mục", "file-system route"),
    ("tuyến giao diện", "web route"),
    ("hàm xử lý tuyến", "route handler"),
    ("điểm cuối API", "API endpoint"),
    ("bộ điều khiển", "controller"),
    ("ca sử dụng", "use case"),
    ("luồng tuần tự", "sequence diagram"),
    ("lược đồ Prisma", "Prisma schema"),
    ("Lược đồ Prisma", "Prisma schema"),
    ("lược đồ Zod", "Zod schema"),
    ("Lược đồ Zod", "Zod schema"),
    ("khóa API", "API key"),
    ("mã truy cập", "access token"),
    ("mã làm mới", "refresh token"),
    ("phiên làm mới", "refresh session"),
    ("giới hạn tần suất", "rate limit"),
    ("mã định danh yêu cầu", "request ID"),
    ("mã băm nội dung", "content hash"),
    ("mã băm", "hash"),
    ("véc-tơ nhúng", "embedding vector"),
    ("véc-tơ", "vector"),
    ("dữ liệu mẫu", "seed data"),
    ("di trú dữ liệu", "data migration"),
    ("di trú", "migration"),
    ("tính lũy đẳng", "idempotency"),
    ("lũy đẳng", "idempotent"),
    ("phương án dự phòng", "fallback"),
    ("độ trễ", "latency"),
    ("thử lại", "retry"),
    ("đánh giá chuẩn", "benchmark"),
    ("mốc tham chiếu", "baseline"),
    ("ảnh chụp", "snapshot"),
    ("mẫu mô phỏng", "mockup"),
    ("trang dữ liệu thô", "raw page"),
    ("dữ liệu đề xuất", "candidate"),
    ("bộ phân tích dữ liệu", "parser"),
    ("k ngữ cảnh phù hợp nhất", "top-k context"),
]


def academic_text(value):
    """Giữ technical terms nguyên bản trong phần diễn giải tiếng Việt."""
    text = str(value)
    for source, target in TECHNICAL_TERM_NORMALIZATION:
        # Dùng biên từ để không làm hỏng thuật ngữ chứa chuỗi con,
        # ví dụ Route không được biến App Router thành "App Tuyếnr".
        if source and source[0].isalnum() and source[-1].isalnum():
            text = re.sub(rf"(?<!\w){re.escape(source)}(?!\w)", target, text)
        else:
            text = text.replace(source, target)
    return text


def load_page_map(path: str | None) -> dict[str, str]:
    if not path:
        return {}
    return json.loads(Path(path).read_text(encoding="utf-8"))


def get_or_add_rpr(run):
    return run._r.get_or_add_rPr()


def set_run_font(run, size=13, bold=None, italic=None, color=BLACK, name="Times New Roman"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rpr = get_or_add_rpr(run)
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), name)


def set_paragraph_font(paragraph, size=13, bold=None, italic=None, color=BLACK):
    for run in paragraph.runs:
        set_run_font(run, size, bold, italic, color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_widths(table, widths_inches):
    table.autofit = False
    for row in table.rows:
        for i, width in enumerate(widths_inches):
            if i >= len(row.cells):
                break
            row.cells[i].width = Inches(width)
            tc_pr = row.cells[i]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths_inches) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")


def add_page_number_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for node in (fld_begin, instr, fld_sep, text, fld_end):
        run._r.append(node)
    set_run_font(run, 11)


def set_page_numbering(section, fmt="decimal", start=1):
    sect_pr = section._sectPr
    node = sect_pr.find(qn("w:pgNumType"))
    if node is None:
        node = OxmlElement("w:pgNumType")
        sect_pr.append(node)
    node.set(qn("w:fmt"), fmt)
    node.set(qn("w:start"), str(start))


def set_page_border(section, color=BLACK, size=18, space=20):
    sect_pr = section._sectPr
    borders = sect_pr.find(qn("w:pgBorders"))
    if borders is None:
        borders = OxmlElement("w:pgBorders")
        borders.set(qn("w:offsetFrom"), "page")
        sect_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), str(space))
        el.set(qn("w:color"), color)
        borders.append(el)


def configure_section(section, numbered=False, fmt="decimal", start=1):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Inches(1.04)
    section.right_margin = Inches(0.77)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.35)
    section.footer.is_linked_to_previous = False
    section.header.is_linked_to_previous = False
    inherited_border = section._sectPr.find(qn("w:pgBorders"))
    if inherited_border is not None:
        section._sectPr.remove(inherited_border)
    for container in (section.footer._element, section.header._element):
        for child in list(container):
            container.remove(child)
        container.append(OxmlElement("w:p"))
    if numbered:
        set_page_numbering(section, fmt, start)
        add_page_number_field(section.footer.paragraphs[0])


def add_marker(paragraph, key):
    if not MARKERS:
        return
    run = paragraph.add_run(f" [[PAGE:{key}]]")
    set_run_font(run, 1, color=WHITE)


def add_body(doc, text="", bold_lead: str | None = None, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             first_line=True, keep=False):
    text = academic_text(text)
    bold_lead = academic_text(bold_lead) if bold_lead else None
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    if first_line:
        p.paragraph_format.first_line_indent = Inches(0.197)
    p.paragraph_format.keep_together = keep
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, 13, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2, 13, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, 13, italic=italic)
    return p


def add_bullets(doc, items):
    num_id = new_numbering_id(doc, "List Bullet")
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        apply_numbering(p, num_id)
        p.paragraph_format.left_indent = Inches(0.55)
        p.paragraph_format.first_line_indent = Inches(-0.30)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.35
        r = p.add_run(academic_text(item))
        set_run_font(r, 13)


def new_numbering_id(doc, style_name="List Number"):
    style = doc.styles[style_name]
    numbering = doc.part.numbering_part.element
    style_p_pr = style._element.pPr
    style_num_pr = style_p_pr.numPr if style_p_pr is not None else None
    style_num_id = style_num_pr.numId if style_num_pr is not None else None
    if style_num_id is None:
        abstract_ids = [
            int(node.get(qn("w:abstractNumId")))
            for node in numbering.findall(qn("w:abstractNum"))
            if (node.get(qn("w:abstractNumId")) or "").isdigit()
        ]
        existing_ids = [
            int(node.get(qn("w:numId")))
            for node in numbering.findall(qn("w:num"))
            if (node.get(qn("w:numId")) or "").isdigit()
        ]
        abstract_id = max(abstract_ids or [0]) + 1
        new_id = max(existing_ids or [0]) + 1
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        level.append(start)
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), "bullet" if style_name == "List Bullet" else "decimal")
        level.append(fmt)
        text = OxmlElement("w:lvlText")
        text.set(qn("w:val"), "•" if style_name == "List Bullet" else "%1.")
        level.append(text)
        suffix = OxmlElement("w:suff")
        suffix.set(qn("w:val"), "tab")
        level.append(suffix)
        p_pr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "864")
        ind.set(qn("w:hanging"), "432")
        p_pr.append(ind)
        level.append(p_pr)
        abstract.append(level)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(new_id))
        ref = OxmlElement("w:abstractNumId")
        ref.set(qn("w:val"), str(abstract_id))
        num.append(ref)
        numbering.append(num)
        return new_id
    base_num_id = str(style_num_id.val)
    abstract_id = None
    existing_ids = []
    for num in numbering.findall(qn("w:num")):
        num_id = num.get(qn("w:numId"))
        if num_id and num_id.isdigit():
            existing_ids.append(int(num_id))
        if num_id == base_num_id:
            ref = num.find(qn("w:abstractNumId"))
            abstract_id = ref.get(qn("w:val")) if ref is not None else None
    if abstract_id is None:
        return int(base_num_id)
    new_id = max(existing_ids or [0]) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), abstract_id)
    num.append(ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return new_id


def apply_numbering(paragraph, num_id, level=0):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.get_or_add_numPr()
    ilvl = num_pr.get_or_add_ilvl()
    ilvl.val = level
    node = num_pr.get_or_add_numId()
    node.val = num_id


def add_numbered(doc, items, compact=False, font_size=None):
    num_id = new_numbering_id(doc)
    for item in items:
        p = doc.add_paragraph(style="List Number")
        apply_numbering(p, num_id)
        p.paragraph_format.left_indent = Inches(0.6)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        p.paragraph_format.space_before = Pt(1 if compact else 2)
        p.paragraph_format.space_after = Pt(1 if compact else 2)
        p.paragraph_format.line_spacing = 1.05 if compact else 1.35
        r = p.add_run(academic_text(item))
        set_run_font(r, font_size if font_size is not None else (11.5 if compact else 13))


def add_heading(doc, text, level, key, toc=True):
    text = academic_text(text)
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    set_run_font(r, 15 if level == 1 else 13, bold=True, italic=(level == 3))
    add_marker(p, key)
    if toc:
        HEADING_ENTRIES.append((level, text, key))
    return p


def add_front_heading(doc, text, key):
    text = academic_text(text)
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.page_break_before = False
    r = p.add_run(text.upper())
    set_run_font(r, 16, bold=True)
    add_marker(p, key)
    HEADING_ENTRIES.append((1, text.upper(), key))
    return p


def add_table_caption(doc, number, text, key):
    text = academic_text(text)
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"Bảng {number}. {text}")
    set_run_font(r, 11, bold=False, italic=True)
    add_marker(p, key)
    TABLE_ENTRIES.append((f"Bảng {number}. {text}", key))
    return p


def add_table(doc, headers, rows, widths=None, font_size=10.5, raw_columns=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, HEADER_FILL)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(academic_text(h))
        set_run_font(r, font_size, bold=True)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for i, value in enumerate(row):
            cell = cells[i]
            if ridx % 2:
                set_cell_shading(cell, ALT_FILL)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.1
            display_value = str(value) if raw_columns and i in raw_columns else academic_text(value)
            r = p.add_run(display_value)
            set_run_font(r, font_size)
    if widths:
        set_table_widths(table, widths)
    return table


def add_figure(doc, path, number, text, key, alt_text, width=6.15, page_break=False):
    text = academic_text(text)
    alt_text = academic_text(alt_text)
    if page_break:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt_text)
    doc_pr.set("title", f"Hình {number}")
    cp = doc.add_paragraph(style="Caption")
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(2)
    cp.paragraph_format.space_after = Pt(8)
    r = cp.add_run(f"Hình {number}. {text}")
    set_run_font(r, 11, italic=True)
    add_marker(cp, key)
    FIGURE_ENTRIES.append((f"Hình {number}. {text}", key))
    return cp


def add_listing(doc, title, code_lines):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.15)
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F2F2F2")
    p_pr.append(shd)
    r = p.add_run(academic_text(title) + "\n" + "\n".join(academic_text(line) for line in code_lines))
    set_run_font(r, 9.5, name="Courier New")


def add_toc_line(doc, level, label, key):
    label = academic_text(label)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.0 if level == 1 else 0.28 if level == 2 else 0.55)
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.05
    p_pr = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), "9000")
    tabs.append(tab)
    p_pr.append(tabs)
    r = p.add_run(label)
    set_run_font(r, 10.5, bold=(level == 1))
    r2 = p.add_run("\t" + PAGE_MAP.get(key, "0"))
    set_run_font(r2, 10.5, bold=(level == 1))


def add_list_line(doc, label, key):
    add_toc_line(doc, 2, label, key)


def extract_logo():
    source = ASSETS / "phenikaa-logo.png"
    out = ASSETS / "phenikaa-logo-bw.png"
    source.parent.mkdir(parents=True, exist_ok=True)
    if not source.exists():
        with ZipFile(TEMPLATE) as z:
            source.write_bytes(z.read("word/media/image1.png"))
    image = Image.open(source).convert("RGBA")
    alpha = image.getchannel("A")
    mono = ImageOps.autocontrast(ImageOps.grayscale(image.convert("RGB")), cutoff=1)
    Image.merge("RGBA", (mono, mono, mono, alpha)).save(out, optimize=True)
    return out


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(13)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Inches(0.197)
    pf.space_before = Pt(5)
    pf.space_after = Pt(5)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for level in range(1, 5):
        style = styles[f"Heading {level}"]
        style_p_pr = style._element.get_or_add_pPr()
        inherited_num_pr = style_p_pr.find(qn("w:numPr"))
        if inherited_num_pr is not None:
            style_p_pr.remove(inherited_num_pr)
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.bold = level <= 3
        style.font.italic = level == 3
        style.font.size = Pt(15 if level == 1 else 13)
        style.paragraph_format.first_line_indent = Inches(0)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.space_before = Pt(0 if level == 1 else 11 if level == 2 else 8)
        style.paragraph_format.space_after = Pt(9 if level == 1 else 5)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
        if level == 1:
            style.paragraph_format.page_break_before = True
    caption = styles["Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    caption.font.size = Pt(11)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor(0, 0, 0)
    for name in ("List Bullet", "List Number"):
        if name not in [style.name for style in styles]:
            styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        styles[name].font.name = "Times New Roman"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        styles[name].font.size = Pt(13)
    if "Code Block" not in [s.name for s in styles]:
        style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Courier New"
        style.font.size = Pt(9.5)


def set_core_properties(doc):
    props = doc.core_properties
    props.title = "Báo cáo đồ án tốt nghiệp – Nền tảng Spechub"
    props.subject = "Phân tích, thiết kế và hiện thực nền tảng tra cứu, so sánh và nghiên cứu thiết bị thông minh"
    props.author = "Nhóm phát triển Spechub"
    props.last_modified_by = "Nhóm phát triển Spechub"
    props.comments = "Báo cáo trình bày quá trình phân tích, thiết kế, xây dựng và đánh giá dự án Spechub."
    props.keywords = "Spechub, catalog, thiết bị thông minh, Next.js, NestJS, PostgreSQL, RAG"
    props.created = datetime(2026, 8, 5, 12, 0, 0)
    props.modified = datetime(2026, 8, 5, 12, 0, 0)


def cover(doc, inner=False):
    logo = extract_logo()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    shape = p.add_run().add_picture(str(logo), width=Inches(1.15))
    shape._inline.docPr.set("descr", "Logo Trường Đại học Phenikaa")
    shape._inline.docPr.set("title", "Logo Phenikaa")
    for text, size, bold, after in [
        ("TRƯỜNG ĐẠI HỌC PHENIKAA", 14, True, 2),
        ("KHOA CÔNG NGHỆ THÔNG TIN", 14, True, 24),
        ("ĐỒ ÁN TỐT NGHIỆP", 20, True, 26),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(after)
        r = p.add_run(text)
        set_run_font(r, size, bold=bold)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    p.paragraph_format.keep_together = True
    r = p.add_run("XÂY DỰNG NỀN TẢNG SPECHUB –\nTRA CỨU, SO SÁNH VÀ NGHIÊN CỨU\nTHIẾT BỊ THÔNG MINH")
    set_run_font(r, 18, bold=True, color=NAVY)
    info = doc.add_table(rows=2 if inner else 1, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.autofit = False
    labels = [("Đơn vị thực hiện:", "Nhóm phát triển Spechub")]
    if inner:
        labels.append(("Loại tài liệu:", "Báo cáo đồ án tốt nghiệp"))
    for i, (a, b) in enumerate(labels):
        info.cell(i, 0).text = a
        info.cell(i, 1).text = b
        for c in info.rows[i].cells:
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p2 in c.paragraphs:
                p2.paragraph_format.first_line_indent = Inches(0)
                p2.paragraph_format.space_after = Pt(3)
                set_paragraph_font(p2, 13, bold=(c is info.cell(i, 0)))
    set_repeat_table_header(info.rows[0])
    set_table_widths(info, [1.8, 3.6])
    tbl_pr = info._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(105 if inner else 120)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("HÀ NỘI – 2026")
    set_run_font(r, 14, bold=True)


def add_summary(doc):
    add_front_heading(doc, "TÓM TẮT", "summary")
    paragraphs = [
        "Spechub là nền tảng dữ liệu và nghiên cứu thiết bị thông minh, được xây dựng nhằm giải quyết tình trạng thông số phân tán, cách đặt tên không thống nhất và thiếu khả năng truy nguyên nguồn khi người dùng tìm hiểu sản phẩm. Hệ thống hợp nhất danh mục thiết bị, biến thể thương mại, linh kiện phần cứng, điểm đánh giá, lịch sử giá và nội dung Wiki trong một mô hình dữ liệu quan hệ có cấu trúc.",
        "Giải pháp sử dụng kiến trúc monorepo với giao diện Next.js 15 và React 19, API NestJS 11 chạy trên Fastify 5, lớp truy cập dữ liệu Prisma 6, PostgreSQL 16 cùng các mở rộng pgvector, pg_trgm và unaccent. Redis được dùng cho phiên đăng nhập, bộ nhớ đệm và điều phối tác vụ; Meilisearch là lựa chọn tăng cường cho tìm kiếm. Các chức năng chính gồm tra cứu và lọc thiết bị, so sánh, khuyến nghị theo nhu cầu, trợ lý AI dựa trên truy xuất tăng cường có trích dẫn, Wiki cộng tác, wishlist, cảnh báo giá, thông báo, quản trị catalog và API B2B.",
        "Báo cáo trình bày quá trình khảo sát yêu cầu, lựa chọn kiến trúc, phân tích ca sử dụng, thiết kế các luồng tương tác, xây dựng quy trình thu thập–duyệt dữ liệu và thiết kế cơ sở dữ liệu theo bốn miền. Kết quả khảo sát mã nguồn ghi nhận 22 tuyến giao diện, 30 bộ điều khiển API với 186 điểm cuối và 139 mô hình Prisma. Trong lần đánh giá ngày 05/08/2026, 36 bộ kiểm thử API gồm 189 ca kiểm thử đều đạt; lược đồ Prisma hợp lệ; PostgreSQL và Redis đáp ứng kiểm tra mức sẵn sàng.",
        "Kết quả cho thấy Spechub đã hình thành được nền tảng kỹ thuật có khả năng mở rộng theo miền nghiệp vụ, duy trì nguồn gốc dữ liệu và hỗ trợ nhiều nhóm người dùng. Các hướng phát triển quan trọng tiếp theo là hoàn thiện đo tải, kiểm thử giao diện đầu cuối, tăng độ phủ dữ liệu thực tế, xây dựng cơ chế đánh giá chất lượng câu trả lời AI và triển khai giám sát sản xuất theo mục tiêu mức dịch vụ.",
    ]
    for text in paragraphs:
        add_body(doc, text)
    add_body(doc, "Từ khóa: Spechub; danh mục thiết bị; so sánh; PostgreSQL; RAG; Wiki; cảnh báo giá.", bold_lead="Từ khóa:", first_line=False)


def add_declaration(doc):
    doc.add_page_break()
    add_front_heading(doc, "LỜI CAM ĐOAN", "declaration")
    for text in [
        "Nhóm thực hiện xin cam đoan báo cáo đồ án tốt nghiệp này là kết quả của quá trình nghiên cứu, phân tích, thiết kế, xây dựng và đánh giá dự án Spechub. Nội dung báo cáo được trình bày trung thực, thống nhất với mã nguồn, lược đồ dữ liệu, cấu hình hệ thống và kết quả kiểm thử tại thời điểm đánh giá.",
        "Các kiến thức, tài liệu và công nghệ được tham khảo đều được ghi rõ trong phần Tài liệu tham khảo. Những số liệu, hình ảnh giao diện và sơ đồ do nhóm thực hiện xây dựng được sử dụng đúng mục đích học thuật. Nhóm thực hiện chịu trách nhiệm về tính chính xác và trung thực của nội dung báo cáo.",
    ]:
        add_body(doc, text)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(50)
    r = p.add_run("Hà Nội, tháng 8 năm 2026\nNhóm phát triển Spechub")
    set_run_font(r, 13, italic=True)


def add_acknowledgements(doc):
    doc.add_page_break()
    add_front_heading(doc, "LỜI CẢM ƠN", "thanks")
    for text in [
        "Nhóm thực hiện xin trân trọng cảm ơn Trường Đại học Phenikaa và Khoa Công nghệ Thông tin đã tạo môi trường học tập, nghiên cứu thuận lợi; đồng thời trang bị những kiến thức cần thiết về kỹ nghệ phần mềm, cơ sở dữ liệu, phát triển ứng dụng web và bảo đảm chất lượng phần mềm.",
        "Nhóm cũng ghi nhận đóng góp của cộng đồng mã nguồn mở đã phát triển Next.js, React, NestJS, Fastify, Prisma, PostgreSQL, Redis, Meilisearch và các công cụ kiểm thử được sử dụng trong dự án. Mọi góp ý về phạm vi, chất lượng dữ liệu, an toàn hệ thống và trải nghiệm người dùng sẽ là cơ sở quan trọng để Spechub tiếp tục hoàn thiện.",
    ]:
        add_body(doc, text)


def toc_entries():
    return [
        (1, "TÓM TẮT", "summary"), (1, "LỜI CAM ĐOAN", "declaration"),
        (1, "LỜI CẢM ƠN", "thanks"), (1, "DANH MỤC HÌNH ẢNH", "lof"),
        (1, "DANH MỤC BẢNG BIỂU", "lot"), (1, "DANH MỤC TỪ VIẾT TẮT", "abbr"),
        (1, "MỞ ĐẦU", "opening"),
        (2, "1. Lý do chọn đề tài", "opening-1"), (2, "2. Mục tiêu", "opening-2"),
        (2, "3. Đối tượng và phạm vi", "opening-3"), (2, "4. Phương pháp thực hiện", "opening-4"),
        (2, "5. Đóng góp của đồ án", "opening-5"), (2, "6. Cấu trúc báo cáo", "opening-6"),
        (1, "CHƯƠNG 1. TỔNG QUAN VỀ ĐỀ TÀI", "ch1"),
        (2, "1.1. Bối cảnh và bài toán", "ch1-1"), (2, "1.2. Khoảng trống của các cách tiếp cận hiện có", "ch1-2"),
        (2, "1.3. Tầm nhìn và phạm vi sản phẩm", "ch1-3"), (2, "1.4. Các bên liên quan", "ch1-4"),
        (2, "1.5. Yêu cầu chức năng", "ch1-5"), (2, "1.6. Yêu cầu phi chức năng", "ch1-6"),
        (2, "1.7. Rủi ro và ràng buộc", "ch1-7"), (2, "1.8. Kết luận chương", "ch1-8"),
        (1, "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ", "ch2"),
        (2, "2.1. Nguyên tắc kiến trúc", "ch2-1"), (2, "2.2. Monorepo và tổ chức mã nguồn", "ch2-2"),
        (2, "2.3. Lớp trình bày với Next.js và React", "ch2-3"), (2, "2.4. API với NestJS và Fastify", "ch2-4"),
        (2, "2.5. Prisma, PostgreSQL và tìm kiếm dữ liệu", "ch2-5"), (2, "2.6. Redis, tác vụ nền và khả năng phục hồi", "ch2-6"),
        (2, "2.7. Truy xuất tăng cường và AI có trích dẫn", "ch2-7"), (2, "2.8. Xác thực, phân quyền và bảo mật", "ch2-8"),
        (2, "2.9. PWA, quan sát hệ thống và triển khai", "ch2-9"), (2, "2.10. Kết luận chương", "ch2-10"),
        (1, "CHƯƠNG 3. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", "ch3"),
        (2, "3.1. Tác nhân và mô hình use case", "ch3-1"), (2, "3.2. Đặc tả use case", "ch3-2"),
        (2, "3.3. Thiết kế sequence diagram", "ch3-3"), (2, "3.4. Thiết kế background workflow", "ch3-4"),
        (2, "3.5. Thiết kế dữ liệu", "ch3-5"), (2, "3.6. Thiết kế API và hợp đồng dữ liệu", "ch3-6"),
        (2, "3.7. Thiết kế an toàn và kiểm soát chất lượng", "ch3-7"), (2, "3.8. Kết luận chương", "ch3-8"),
        (1, "CHƯƠNG 4. PHÁT TRIỂN VÀ KẾT QUẢ", "ch4"),
        (2, "4.1. Môi trường và quy trình phát triển", "ch4-1"), (2, "4.2. Hiện thực các mô-đun chính", "ch4-2"),
        (2, "4.3. Kết quả giao diện", "ch4-3"), (2, "4.4. Kết quả API và dữ liệu", "ch4-4"),
        (2, "4.5. Kiểm thử và đánh giá", "ch4-5"), (2, "4.6. Đánh giá mức độ đáp ứng", "ch4-6"),
        (2, "4.7. Hạn chế", "ch4-7"), (2, "4.8. Kết luận chương", "ch4-8"),
        (1, "KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", "conclusion"),
        (1, "TÀI LIỆU THAM KHẢO", "references"),
        (1, "PHỤ LỤC A. DANH MỤC API", "appendix-a"),
        (1, "PHỤ LỤC B. DANH MỤC MÔ HÌNH DỮ LIỆU", "appendix-b"),
        (1, "PHỤ LỤC C. HƯỚNG DẪN VẬN HÀNH VÀ KIỂM THỬ", "appendix-c"),
        (1, "PHỤ LỤC D. HƯỚNG DẪN SỬ DỤNG TÓM TẮT", "appendix-d"),
    ]


def add_toc(doc):
    doc.add_page_break()
    add_front_heading(doc, "MỤC LỤC", "toc")
    entries = toc_entries()
    for i, entry in enumerate(entries):
        if i == 28:
            doc.add_page_break()
            add_front_heading(doc, "MỤC LỤC (TIẾP)", f"toc-cont-{i}")
        add_toc_line(doc, *entry)


def figure_catalog():
    return [
        ("Hình 1.1. Sơ đồ ngữ cảnh hệ thống Spechub", "fig-1-1"),
        ("Hình 2.1. Cấu trúc monorepo Spechub", "fig-2-1"),
        ("Hình 2.2. Kiến trúc logic nhiều lớp", "fig-2-2"),
        ("Hình 2.3. Mô hình triển khai tham chiếu", "fig-2-3"),
        ("Hình 3.1. Use case phía người dùng", "fig-3-1"),
        ("Hình 3.2. Use case biên tập và quản trị", "fig-3-2"),
        ("Hình 3.3. Sequence diagram đăng nhập và refresh session", "fig-3-3"),
        ("Hình 3.4. Sequence diagram tìm kiếm và so sánh", "fig-3-4"),
        ("Hình 3.5. Sequence diagram trợ lý AI theo RAG", "fig-3-5"),
        ("Hình 3.6. Workflow theo dõi giá và phát cảnh báo", "fig-3-6"),
        ("Hình 3.7. Ingestion workflow", "fig-3-7"),
        ("Hình 3.8. ERD nhóm catalog", "fig-3-8"),
        ("Hình 3.9. ERD nhóm hardware và scoring", "fig-3-9"),
        ("Hình 3.10. ERD nhóm content, crawler và AI/search", "fig-3-10"),
        ("Hình 3.11. ERD nhóm user engagement, subscription và B2B", "fig-3-11"),
        ("Hình 4.1. Trang chủ Spechub", "fig-4-1"),
        ("Hình 4.2. Trang danh mục thiết bị", "fig-4-2"),
        ("Hình 4.3. Trang chi tiết thiết bị", "fig-4-3"),
        ("Hình 4.4. Kết quả tìm kiếm thiết bị và phần cứng", "fig-4-4"),
        ("Hình 4.5. Giao diện chọn thiết bị để so sánh", "fig-4-5"),
        ("Hình 4.6. Quy trình khuyến nghị theo nhu cầu", "fig-4-6"),
        ("Hình 4.7. Giao diện trợ lý AI", "fig-4-7"),
        ("Hình 4.8. Giao diện Wiki công nghệ", "fig-4-8"),
        ("Hình 4.9. Giao diện đăng nhập", "fig-4-9"),
    ]


def table_catalog():
    return [
        ("Bảng 1.1. So sánh cách tiếp cận dữ liệu thiết bị", "tab-1-1"),
        ("Bảng 1.2. Các bên liên quan", "tab-1-2"),
        ("Bảng 1.3. Yêu cầu chức năng chính", "tab-1-3"),
        ("Bảng 1.4. Yêu cầu phi chức năng", "tab-1-4"),
        ("Bảng 1.5. Rủi ro và biện pháp giảm thiểu", "tab-1-5"),
        ("Bảng 2.1. Thành phần công nghệ và phiên bản", "tab-2-1"),
        ("Bảng 2.2. Biện pháp bảo mật theo lớp", "tab-2-2"),
        ("Bảng 2.3. Lựa chọn kiến trúc và đánh đổi", "tab-2-3"),
        ("Bảng 3.1. Tác nhân hệ thống", "tab-3-1"),
        ("Bảng 3.2. Danh mục use case", "tab-3-2"),
        ("Bảng 3.3. Đặc tả UC-01 – Tìm kiếm và lọc thiết bị", "tab-3-3"),
        ("Bảng 3.4. Đặc tả UC-02 – So sánh thiết bị", "tab-3-4"),
        ("Bảng 3.5. Đặc tả UC-03 – Nhận khuyến nghị và hỏi AI", "tab-3-5"),
        ("Bảng 3.6. Đặc tả UC-04 – Đăng nhập và quản lý phiên", "tab-3-6"),
        ("Bảng 3.7. Đặc tả UC-05 – Wishlist và cảnh báo giá", "tab-3-7"),
        ("Bảng 3.8. Đặc tả UC-06 – Đóng góp và duyệt Wiki", "tab-3-8"),
        ("Bảng 3.9. Đặc tả UC-07 – Thu thập và duyệt catalog", "tab-3-9"),
        ("Bảng 3.10. Đặc tả UC-08 – Khai thác API B2B", "tab-3-10"),
        ("Bảng 3.11. Phân nhóm 139 mô hình dữ liệu", "tab-3-11"),
        ("Bảng 3.12. Quy ước hợp đồng API", "tab-3-12"),
        ("Bảng 4.1. Môi trường phần mềm", "tab-4-1"),
        ("Bảng 4.2. Mô-đun hiện thực chính", "tab-4-2"),
        ("Bảng 4.3. Bản đồ route giao diện", "tab-4-3"),
        ("Bảng 4.4. Thống kê endpoint API", "tab-4-4"),
        ("Bảng 4.5. Bằng chứng kiểm thử và kiểm chứng", "tab-4-5"),
        ("Bảng 4.6. Đánh giá mức độ đáp ứng yêu cầu", "tab-4-6"),
        ("Bảng 4.7. Hạn chế và tác động", "tab-4-7"),
        ("Bảng A.1. Danh mục controller và endpoint", "tab-a-1"),
        ("Bảng B.1. Danh sách mô hình Prisma", "tab-b-1"),
        ("Bảng C.1. Ma trận lệnh vận hành và mục đích", "tab-c-1"),
        ("Bảng D.1. Tác vụ người dùng thường gặp", "tab-d-1"),
    ]


def add_front_lists(doc):
    doc.add_page_break()
    add_front_heading(doc, "DANH MỤC HÌNH ẢNH", "lof")
    for i, (label, key) in enumerate(figure_catalog()):
        if i == 13:
            doc.add_page_break()
            add_front_heading(doc, "DANH MỤC HÌNH ẢNH (TIẾP)", "lof-cont")
        add_list_line(doc, label, key)
    doc.add_page_break()
    add_front_heading(doc, "DANH MỤC BẢNG BIỂU", "lot")
    for i, (label, key) in enumerate(table_catalog()):
        if i == 16:
            doc.add_page_break()
            add_front_heading(doc, "DANH MỤC BẢNG BIỂU (TIẾP)", "lot-cont")
        add_list_line(doc, label, key)
    doc.add_page_break()
    add_front_heading(doc, "DANH MỤC TỪ VIẾT TẮT", "abbr")
    rows = [
        ("AI", "Artificial Intelligence", "Trí tuệ nhân tạo"),
        ("API", "Application Programming Interface", "Giao diện lập trình ứng dụng"),
        ("B2B", "Business-to-Business", "Dịch vụ giữa các tổ chức"),
        ("CI", "Continuous Integration", "Tích hợp liên tục"),
        ("CRUD", "Create, Read, Update, Delete", "Các thao tác dữ liệu cơ bản"),
        ("CSR", "Client-Side Rendering", "Kết xuất phía trình duyệt"),
        ("DTO", "Data Transfer Object", "Đối tượng truyền dữ liệu"),
        ("ERD", "Entity–Relationship Diagram", "Sơ đồ thực thể–quan hệ"),
        ("FK", "Foreign Key", "Khóa ngoại"),
        ("HTTP(S)", "Hypertext Transfer Protocol (Secure)", "Giao thức truyền siêu văn bản (bảo mật)"),
        ("JWT", "JSON Web Token", "Mã thông báo web JSON"),
        ("LLM", "Large Language Model", "Mô hình ngôn ngữ lớn"),
        ("NFR", "Non-functional Requirement", "Yêu cầu phi chức năng"),
        ("ORM", "Object–Relational Mapping", "Ánh xạ đối tượng–quan hệ"),
        ("PK", "Primary Key", "Khóa chính"),
        ("PWA", "Progressive Web Application", "Ứng dụng web tiến bộ"),
        ("RAG", "Retrieval-Augmented Generation", "Sinh nội dung có truy xuất tăng cường"),
        ("RBAC", "Role-Based Access Control", "Kiểm soát truy cập theo vai trò"),
        ("REST", "Representational State Transfer", "Phong cách kiến trúc dịch vụ web"),
        ("SDK", "Software Development Kit", "Bộ công cụ phát triển phần mềm"),
        ("SLO", "Service Level Objective", "Mục tiêu mức dịch vụ"),
        ("SPA", "Single-Page Application", "Ứng dụng đơn trang"),
        ("SSR", "Server-Side Rendering", "Kết xuất phía máy chủ"),
        ("UI/UX", "User Interface / User Experience", "Giao diện / trải nghiệm người dùng"),
        ("UUID", "Universally Unique Identifier", "Định danh duy nhất phổ quát"),
    ]
    add_table(doc, ["Từ viết tắt", "Tiếng Anh", "Ý nghĩa"], rows, [1.0, 2.8, 2.4], 10.5)


def add_opening(doc):
    add_heading(doc, "MỞ ĐẦU", 1, "opening")
    add_heading(doc, "1. Lý do chọn đề tài", 2, "opening-1")
    for text in [
        "Thị trường thiết bị thông minh phát triển nhanh về số lượng model, biến thể bộ nhớ, cấu hình phần cứng và phiên bản phần mềm. Cùng một sản phẩm có thể được mô tả bằng nhiều tên thương mại, đơn vị đo hoặc tập thông số khác nhau giữa trang hãng, cửa hàng và bài đánh giá. Người dùng phải mở nhiều nguồn, tự đối chiếu và thường không biết một con số bắt nguồn từ đâu.",
        "Các trang thương mại điện tử tối ưu cho giao dịch, trong khi các bảng thông số đơn lẻ thường thiếu quan hệ giữa thiết bị, chipset, CPU, GPU, màn hình, camera và bằng chứng nguồn. Khi bổ sung AI, nếu hệ thống không ràng buộc câu trả lời vào catalog đã chuẩn hóa và trích dẫn, nguy cơ đưa ra thông tin thiếu căn cứ càng tăng. Spechub được lựa chọn để nghiên cứu một nền tảng coi dữ liệu, nguồn gốc và quy trình duyệt là lõi, thay vì chỉ là một giao diện tìm kiếm.",
    ]:
        add_body(doc, text)
    add_heading(doc, "2. Mục tiêu", 2, "opening-2")
    add_body(doc, "Mục tiêu tổng quát là xây dựng nền tảng web phục vụ tra cứu, so sánh và nghiên cứu thiết bị thông minh trên một mô hình dữ liệu có cấu trúc, có thể truy nguyên nguồn và mở rộng cho tác vụ AI. Các mục tiêu cụ thể gồm:")
    add_bullets(doc, [
        "Chuẩn hóa danh mục model, variant, linh kiện và thuộc tính kỹ thuật.",
        "Cung cấp tìm kiếm, lọc, xem chi tiết, so sánh và khuyến nghị theo nhu cầu.",
        "Thiết kế trợ lý AI/RAG sử dụng ngữ cảnh catalog và trả về trích dẫn.",
        "Tổ chức Wiki có phiên bản, kiểm duyệt, bình luận và liên kết nguồn.",
        "Hỗ trợ wishlist, lịch sử giá, cảnh báo và thông báo theo kênh.",
        "Cung cấp công cụ biên tập, nhật ký hoạt động, kiểm tra trạng thái, chỉ số giám sát và API dành cho đối tác có hạn mức.",
        "Đánh giá hiện thực bằng kiểm thử tự động, kiểm tra lược đồ và readiness hạ tầng.",
    ])
    add_heading(doc, "3. Đối tượng và phạm vi", 2, "opening-3")
    add_body(doc, "Đối tượng nghiên cứu gồm kiến trúc ứng dụng web nhiều lớp, thiết kế cơ sở dữ liệu danh mục, tìm kiếm toàn văn và véc-tơ, quản lý phiên và phân quyền, quy trình biên tập dữ liệu, RAG có trích dẫn, tác vụ nền và kiểm thử API. Phạm vi thực hiện bao gồm giao diện web, API, tiến trình nền, các gói dùng chung và hạ tầng PostgreSQL/Redis. Việc đánh giá tải lớn trên môi trường Internet và mức độ bao phủ toàn bộ thiết bị trên thị trường không nằm trong phạm vi của đồ án.")
    add_heading(doc, "4. Phương pháp thực hiện", 2, "opening-4")
    add_numbered(doc, [
        "Khảo sát bài toán, xác định đối tượng sử dụng, phạm vi nghiên cứu và các yêu cầu của hệ thống.",
        "Phân tích tài liệu dự án, cấu trúc mã nguồn, các bộ điều khiển API, tuyến giao diện và lược đồ Prisma.",
        "Vận hành hệ thống trong môi trường cục bộ, kiểm tra mức sẵn sàng, ghi nhận giao diện và thực hiện các bộ kiểm thử API.",
        "Mô hình hóa tác nhân, ca sử dụng, biểu đồ tuần tự, quy trình nghiệp vụ và sơ đồ thực thể–liên kết theo từng miền dữ liệu.",
        "Đối chiếu kết quả xây dựng với yêu cầu, phân tích các hạn chế và đề xuất hướng phát triển tiếp theo.",
    ], compact=True)
    add_heading(doc, "5. Đóng góp của đồ án", 2, "opening-5")
    add_body(doc, "Đồ án đóng góp một mô hình dữ liệu sâu cho thiết bị thông minh; một kiến trúc monorepo tách ứng dụng triển khai và gói dùng chung; một bề mặt API có phân quyền; luồng AI gắn với dữ liệu và trích dẫn; cùng quy trình biên tập không cho dữ liệu ngoài đi thẳng vào catalog công khai. Báo cáo cũng đưa ra bộ sơ đồ và phụ lục định lượng để các thành viên mới có thể hiểu nhanh hệ thống.")
    add_heading(doc, "6. Cấu trúc báo cáo", 2, "opening-6")
    add_body(doc, "Sau phần mở đầu, Chương 1 trình bày bối cảnh, yêu cầu và ràng buộc; Chương 2 giải thích cơ sở lý thuyết và các lựa chọn công nghệ; Chương 3 đặc tả use case, luồng tương tác và dữ liệu; Chương 4 mô tả hiện thực, giao diện và kết quả kiểm thử. Phần cuối tổng kết, nêu hướng phát triển, tài liệu tham khảo và các phụ lục API, mô hình dữ liệu, vận hành và hướng dẫn sử dụng.")


def add_chapter_1(doc):
    add_heading(doc, "CHƯƠNG 1. TỔNG QUAN VỀ ĐỀ TÀI", 1, "ch1")
    add_heading(doc, "1.1. Bối cảnh và bài toán", 2, "ch1-1")
    for text in [
        "Một quyết định mua thiết bị không chỉ phụ thuộc giá niêm yết. Người dùng cần biết khác biệt giữa các variant, cấu hình chipset, chất lượng màn hình, camera, pin, khả năng cập nhật hệ điều hành và mức phù hợp với nhu cầu. Dữ liệu này thường nằm rải rác ở trang hãng, cửa hàng, cơ sở benchmark và bài viết kỹ thuật. Tên trường, đơn vị và mức chi tiết khác nhau khiến việc tổng hợp thủ công tốn thời gian và dễ sai.",
        "Ở góc độ kỹ thuật, dữ liệu thiết bị có tính quan hệ cao. Một model có nhiều variant; một variant dùng nhiều thành phần; một linh kiện có thể xuất hiện trên nhiều thiết bị; một giá trị có thể cần nhiều nguồn xác nhận và thay đổi theo thời gian. Nếu lưu toàn bộ dưới dạng văn bản hoặc JSON không được quản trị, việc tìm kiếm, so sánh, chấm điểm và duy trì tính nhất quán sẽ khó mở rộng.",
        "Spechub xác định ba bài toán trung tâm: chuẩn hóa tri thức thiết bị; cung cấp trải nghiệm khám phá có thể giải thích; và duy trì vòng đời dữ liệu từ thu thập, đối soát, xuất bản đến lập chỉ mục. Các chức năng thương mại như wishlist, cảnh báo giá, affiliate và gói API được thiết kế xoay quanh lõi dữ liệu đó.",
    ]:
        add_body(doc, text)
    add_figure(doc, DIAGRAMS / "01-context.png", "1.1", "Sơ đồ ngữ cảnh hệ thống Spechub", "fig-1-1",
               "Sơ đồ ngữ cảnh cho thấy khách truy cập, người dùng, biên tập viên, quản trị viên và đối tác B2B tương tác với nền tảng Spechub.")

    add_heading(doc, "1.2. Khoảng trống của các cách tiếp cận hiện có", 2, "ch1-2")
    add_body(doc, "Không có một kiểu sản phẩm đơn lẻ đáp ứng đồng thời chiều sâu dữ liệu, trải nghiệm so sánh, cộng tác nội dung, khả năng truy nguyên và quyền truy cập lập trình. Bảng 1.1 tổng hợp sự khác biệt theo mục tiêu sử dụng; đây là phân tích định tính về loại hình giải pháp, không phải xếp hạng một nhà cung cấp cụ thể.")
    add_table_caption(doc, "1.1", "So sánh cách tiếp cận dữ liệu thiết bị", "tab-1-1")
    add_table(doc,
              ["Cách tiếp cận", "Điểm mạnh", "Khoảng trống đối với Spechub"],
              [
                  ("Trang hãng", "Thông tin chính thức, hình ảnh chuẩn", "Khó so sánh chéo; dữ liệu thường chỉ tập trung sản phẩm của hãng"),
                  ("Sàn/cửa hàng", "Giá và tồn kho gần giao dịch", "Thông số có thể rút gọn, lặp hoặc không đồng nhất giữa người bán"),
                  ("Trang thông số", "Tra cứu nhanh nhiều model", "Quan hệ linh kiện, nguồn trích dẫn và quy trình duyệt thường hạn chế"),
                  ("Bài đánh giá", "Phân tích sâu, có ngữ cảnh sử dụng", "Khó cấu trúc hóa để lọc, so sánh tự động và cập nhật hàng loạt"),
                  ("Chatbot tổng quát", "Hội thoại tự nhiên", "Có nguy cơ trả lời không bám catalog, thiếu trích dẫn và không tái lập"),
                  ("Spechub", "Kết hợp catalog, nguồn, Wiki, so sánh, RAG và cảnh báo", "Đòi hỏi đầu tư lớn cho chuẩn hóa, kiểm duyệt và vận hành dữ liệu"),
              ], [1.25, 2.35, 2.6])
    add_body(doc, "Khoảng trống mà Spechub ưu tiên không phải là số lượng trang, mà là khả năng liên kết dữ liệu có cấu trúc với bằng chứng. Khi một thuộc tính xuất hiện trong kết quả so sánh hoặc câu trả lời AI, hệ thống cần xác định được thực thể liên quan, phiên bản dữ liệu và nguồn tham chiếu phù hợp.")

    add_heading(doc, "1.3. Tầm nhìn và phạm vi sản phẩm", 2, "ch1-3")
    add_body(doc, "Tầm nhìn của Spechub là trở thành lớp tri thức dùng chung cho nghiên cứu thiết bị thông minh. Lớp trải nghiệm phục vụ người dùng cuối; Catalog Studio phục vụ quản trị dữ liệu; API B2B mở quyền truy cập có kiểm soát; còn worker duy trì các tác vụ dài và định kỳ. Phạm vi mã nguồn quan sát được bao gồm 22 route web, 30 controller API, 186 endpoint và 139 mô hình dữ liệu.")
    add_bullets(doc, [
        "Catalog: loại thiết bị, hãng/tổ chức, họ sản phẩm, model, variant, alias, media và trạng thái phát hành.",
        "Phần cứng: chipset, CPU, GPU, NPU, màn hình, camera, pin, bộ nhớ, lưu trữ, hệ điều hành và benchmark.",
        "Khám phá: tìm kiếm, lọc, xem chi tiết, so sánh, chấm điểm và khuyến nghị.",
        "Nội dung: nguồn, trích dẫn, Wiki, phiên bản, bình luận, kiểm duyệt và embedding.",
        "Cá nhân hóa: tài khoản, wishlist, cảnh báo giá, thông báo và tùy chọn kênh nhận.",
        "Thương mại và vận hành: tiếp thị liên kết, lịch sử giá, gói thuê bao, Stripe webhook, khóa API, thống kê sử dụng, nhật ký hoạt động, kiểm tra trạng thái và chỉ số giám sát.",
    ])

    add_heading(doc, "1.4. Các bên liên quan", 2, "ch1-4")
    add_table_caption(doc, "1.2", "Các bên liên quan", "tab-1-2")
    add_table(doc,
              ["Bên liên quan", "Nhu cầu chính", "Tiêu chí thành công"],
              [
                  ("Khách truy cập", "Tìm và so sánh nhanh mà không buộc đăng nhập", "Kết quả rõ, có bộ lọc, có nguồn và URL chia sẻ được"),
                  ("Người dùng", "Lưu danh sách, theo dõi giá, nhận thông báo", "Dữ liệu cá nhân tách biệt, cảnh báo đúng điều kiện, không gửi lặp"),
                  ("Biên tập viên", "Cập nhật catalog và nội dung từ nguồn", "Có hàng đợi duyệt, diff, citation và audit"),
                  ("Kiểm duyệt viên", "Đánh giá thay đổi và phiên bản Wiki", "Phê duyệt/từ chối có lý do, phân quyền đúng"),
                  ("Quản trị viên", "Quản lý quyền, gói dịch vụ và vận hành", "Có dashboard, health, metrics, log và cơ chế thu hồi"),
                  ("Đối tác B2B", "Truy cập dữ liệu ổn định bằng API", "API key có scope, hạn mức, thống kê usage và tài liệu hợp đồng"),
                  ("Nhóm phát triển", "Phát triển nhiều mô-đun trong một kho", "Build/test nhất quán, kiểu dữ liệu dùng chung, lỗi dễ truy vết"),
              ], [1.35, 2.45, 2.4])

    add_heading(doc, "1.5. Yêu cầu chức năng", 2, "ch1-5")
    add_body(doc, "Yêu cầu chức năng được nhóm theo hành trình người dùng và vòng đời dữ liệu. Mỗi yêu cầu có mã để đối chiếu ở Chương 4.")
    add_table_caption(doc, "1.3", "Yêu cầu chức năng chính", "tab-1-3")
    add_table(doc,
              ["Mã", "Yêu cầu", "Mức ưu tiên", "Bằng chứng hiện thực"],
              [
                  ("FR-01", "Tìm kiếm thiết bị/phần cứng theo từ khóa và bộ lọc", "Bắt buộc", "Route /search, controller search/catalog"),
                  ("FR-02", "Xem model, variant, thông số, điểm và nguồn", "Bắt buộc", "Route /devices/[slug], API catalog"),
                  ("FR-03", "So sánh nhiều thiết bị trên thuộc tính chuẩn hóa", "Bắt buộc", "Route /compare, compare service"),
                  ("FR-04", "Đề xuất thiết bị theo nhu cầu", "Cao", "Route /recommend, AI recommendations"),
                  ("FR-05", "Hỏi AI và nhận câu trả lời có citation", "Cao", "Route /ai, /ai/ask và /ai/chat"),
                  ("FR-06", "Đăng ký, đăng nhập, làm mới và thu hồi phiên", "Bắt buộc", "Auth API, JWT và Redis session"),
                  ("FR-07", "Quản lý wishlist và cảnh báo giá", "Cao", "Route /wishlist, /alerts và worker"),
                  ("FR-08", "Đọc, soạn, sửa và duyệt Wiki theo phiên bản", "Cao", "Các route /wiki và API moderation"),
                  ("FR-09", "Thu thập, chuẩn hóa và duyệt dữ liệu danh mục", "Bắt buộc", "Nguồn dữ liệu, trang dữ liệu thô và quy trình kiểm duyệt"),
                  ("FR-10", "Quản trị quyền, gói thuê bao, nhật ký hoạt động và webhook", "Cao", "Các bộ điều khiển quản trị, thanh toán và nhật ký hoạt động"),
                  ("FR-11", "Cấp khóa API, phạm vi quyền và theo dõi mức sử dụng B2B", "Trung bình", "Tuyến /api-access và các bộ điều khiển B2B"),
                  ("FR-12", "Cung cấp kiểm tra trạng thái, mức sẵn sàng và chỉ số giám sát", "Bắt buộc", "Bốn điểm cuối kiểm tra trạng thái hoạt động"),
              ], [0.7, 2.95, 0.9, 1.65], 9.8)

    add_heading(doc, "1.6. Yêu cầu phi chức năng", 2, "ch1-6")
    add_table_caption(doc, "1.4", "Yêu cầu phi chức năng", "tab-1-4")
    add_table(doc,
              ["Mã", "Thuộc tính", "Yêu cầu có thể kiểm chứng"],
              [
                  ("NFR-01", "An toàn", "Validate DTO theo whitelist; từ chối trường lạ; JWT và RBAC; Helmet; CORS giới hạn; không lưu API key dạng rõ."),
                  ("NFR-02", "Tính đúng", "Ràng buộc khóa/unique trong DB; quy trình duyệt; citation; kiểm thử service/controller; schema Prisma hợp lệ."),
                  ("NFR-03", "Khả dụng", "Health/liveness/readiness tách biệt; phụ thuộc DB/Redis được báo rõ; lỗi có request ID."),
                  ("NFR-04", "Hiệu năng", "Phân trang, index, cache Redis, tìm kiếm chuyên dụng tùy chọn; tác vụ dài chạy nền."),
                  ("NFR-05", "Mở rộng", "Mô-đun NestJS theo miền; monorepo; dữ liệu chuẩn hóa; API version /v1."),
                  ("NFR-06", "Quan sát", "Structured log, request ID, metrics và audit cho hành động đặc quyền."),
                  ("NFR-07", "Trải nghiệm", "Giao diện responsive, PWA/offline, thông báo trạng thái và thao tác chính có thể dùng bàn phím."),
                  ("NFR-08", "Bảo trì", "TypeScript strict, schema dùng chung, lint/test, tài liệu Swagger trong môi trường phát triển."),
              ], [0.75, 1.05, 4.4], 10.0)
    add_body(doc, "Một số mục như ngưỡng độ trễ theo phân vị, tải đồng thời và mục tiêu mức dịch vụ trong môi trường vận hành chưa có kết quả đo tải trong phạm vi đồ án. Vì vậy, các nội dung này được trình bày dưới dạng mục tiêu thiết kế và hướng đánh giá tiếp theo, không sử dụng số liệu hiệu năng giả định.")

    add_heading(doc, "1.7. Rủi ro và ràng buộc", 2, "ch1-7")
    add_table_caption(doc, "1.5", "Rủi ro và biện pháp giảm thiểu", "tab-1-5")
    add_table(doc,
              ["Rủi ro", "Khả năng / tác động", "Biện pháp giảm thiểu"],
              [
                  ("Nguồn ngoài thay đổi cấu trúc", "Cao / Cao", "Lưu raw page, content hash, parser theo nguồn, hàng đợi review và cảnh báo thất bại"),
                  ("Thông số xung đột", "Cao / Cao", "Chuẩn hóa đơn vị, lưu citation, trust score và yêu cầu đối soát trước xuất bản"),
                  ("AI tạo nội dung thiếu căn cứ", "Trung bình / Cao", "RAG top-k, prompt policy, citation ID, cache có phiên bản và đánh giá thủ công"),
                  ("Lộ refresh token/API key", "Thấp / Cao", "Cookie bảo mật, hash key, session Redis, scope, rate limit và thu hồi"),
                  ("Tăng trưởng lược đồ", "Cao / Trung bình", "Chia miền, bảng nối, migration, quy ước đặt tên và phụ lục mô hình"),
                  ("Phụ thuộc dịch vụ ngoài", "Trung bình / Trung bình", "Timeout, fallback, retry có giới hạn, webhook idempotency và degrade có kiểm soát"),
                  ("Thiếu dữ liệu bao phủ", "Cao / Trung bình", "Ưu tiên chất lượng hơn số lượng, đo độ đầy đủ theo category và mở rộng theo lộ trình"),
              ], [1.65, 1.4, 3.15], 10.0)

    add_heading(doc, "1.8. Kết luận chương", 2, "ch1-8")
    add_body(doc, "Chương 1 đã xác định Spechub là bài toán quản trị tri thức thiết bị chứ không chỉ là một website liệt kê sản phẩm. Yêu cầu trọng tâm là cấu trúc dữ liệu sâu, truy nguyên nguồn, khám phá dễ dùng, kiểm soát thay đổi và bề mặt API an toàn. Các yêu cầu và rủi ro này là cơ sở để lựa chọn công nghệ ở Chương 2 và đặc tả thiết kế ở Chương 3.")


def add_chapter_2(doc, evidence):
    add_heading(doc, "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ", 1, "ch2")
    add_heading(doc, "2.1. Nguyên tắc kiến trúc", 2, "ch2-1")
    for text in [
        "Kiến trúc Spechub tuân theo các nguyên tắc: phân tách trách nhiệm; hợp đồng dữ liệu rõ; xác thực ở biên và phân quyền ở nghiệp vụ; mọi thay đổi quan trọng có audit; dữ liệu ngoài phải qua bước chuẩn hóa/duyệt; tác vụ dài được tách khỏi request; và chức năng AI phải có khả năng truy vết về ngữ cảnh.",
        "Ứng dụng được tổ chức theo nhiều lớp. Web chịu trách nhiệm trình bày và trạng thái tương tác; API thực hiện validation, xác thực và điều phối use case; service/module chứa quy tắc nghiệp vụ; Prisma quản lý truy cập quan hệ; còn PostgreSQL, Redis, search và các dịch vụ ngoài đảm nhiệm lưu trữ hoặc năng lực chuyên biệt. Cách phân tách này giúp kiểm thử từng tầng và thay thế phụ thuộc mà không làm thay đổi toàn bộ giao diện.",
    ]:
        add_body(doc, text)
    add_table_caption(doc, "2.1", "Thành phần công nghệ và phiên bản", "tab-2-1")
    versions = evidence["versions"]
    rows = [
        ("Môi trường thực thi", "Node.js", versions.get("node", ""), "Thực thi web/API và công cụ"),
        ("Quản lý kho mã nguồn", "pnpm / Turborepo", f"{versions.get('packageManager', '')} / {versions.get('turbo', '')}", "Quản lý kho mã nguồn đơn nhất và quy trình xử lý"),
        ("Giao diện web", "Next.js / React", f"{versions.get('next', '')} / {versions.get('react', '')}", "SSR/CSR, App Router và giao diện"),
        ("API", "NestJS / Fastify", f"{versions.get('@nestjs/core', '')} / {versions.get('fastify', '')}", "REST, guard, pipe và hiệu năng HTTP"),
        ("Kiểm tra dữ liệu", "Zod", versions.get("zod", ""), "Lược đồ dùng chung và kiểm tra dữ liệu"),
        ("Truy cập dữ liệu", "Prisma", versions.get("prisma", ""), "ORM, di trú và kiểu dữ liệu"),
        ("Cơ sở dữ liệu", "PostgreSQL", "16", "Quan hệ, tìm kiếm toàn văn, véc-tơ và phần mở rộng"),
        ("Bộ nhớ đệm/phiên", "Redis", "Theo môi trường triển khai", "Phiên, bộ nhớ đệm và điều phối"),
        ("Tìm kiếm", "Meilisearch client", versions.get("meilisearch", ""), "Tìm kiếm chuyên dụng tùy chọn"),
        ("Trạng thái phía web", "TanStack Query", versions.get("@tanstack/react-query", ""), "Bộ nhớ đệm và đồng bộ dữ liệu phía web"),
    ]
    add_table(doc, ["Lớp", "Công nghệ", "Phiên bản", "Vai trò"], rows, [1.05, 1.55, 1.65, 1.95], 9.7)

    add_heading(doc, "2.2. Monorepo và tổ chức mã nguồn", 2, "ch2-2")
    add_body(doc, "Monorepo dùng pnpm workspace và Turborepo cho phép web, API, worker, database và các gói dùng chung phát triển trong một kho nhưng vẫn có ranh giới triển khai. Kiểu dữ liệu, schema validation và cấu hình TypeScript/ESLint được chia sẻ để giảm sai khác giữa client và server. Pipeline có thể chạy build, test và lint theo đồ thị phụ thuộc, tận dụng cache cho phần không thay đổi.")
    add_figure(doc, DIAGRAMS / "15-monorepo.png", "2.1", "Cấu trúc monorepo Spechub", "fig-2-1",
               "Sơ đồ cấu trúc monorepo gồm các ứng dụng web, API, worker và các gói database, shared, config.")
    add_body(doc, "Ranh giới package không thay thế ranh giới nghiệp vụ. Trong API, mỗi miền như catalog, AI, Wiki, alerts, affiliate, billing và admin tiếp tục có module/controller/service riêng. Cách tổ chức hai cấp giúp đội phát triển xác định nơi chứa code, đồng thời hạn chế việc một module truy cập trực tiếp chi tiết nội bộ của module khác.")

    add_heading(doc, "2.3. Lớp trình bày với Next.js và React", 2, "ch2-3")
    for text in [
        "Next.js App Router tổ chức trang theo route filesystem, hỗ trợ kết xuất phía máy chủ cho nội dung cần khả năng lập chỉ mục và kết xuất phía khách cho tương tác giàu trạng thái. React 19 cung cấp mô hình component; TanStack Query quản lý cache, trạng thái tải/lỗi và làm mới dữ liệu server; Tailwind CSS giúp thống nhất token khoảng cách, màu và responsive.",
        "Danh sách 22 route thể hiện bốn nhóm trải nghiệm: khám phá công khai; tài khoản/cá nhân hóa; nội dung Wiki; và quản trị/thương mại. Thiết kế PWA bổ sung service worker và trang offline để người dùng có phản hồi rõ khi mất kết nối. Việc phân tách route không đồng nghĩa mọi trang đều công khai: middleware và trạng thái xác thực phải điều hướng hoặc chặn thao tác cần quyền.",
    ]:
        add_body(doc, text)
    add_bullets(doc, [
        "Server rendering phù hợp trang catalog/chi tiết cần metadata và tải lần đầu ổn định.",
        "Client component dùng cho bộ lọc, chọn so sánh, biểu mẫu nhiều bước và hội thoại AI.",
        "TanStack Query giúp tránh tự viết cơ chế cache riêng cho từng trang.",
        "URL chứa slug và query filter làm trạng thái có thể chia sẻ, quay lại và kiểm thử.",
    ])

    add_heading(doc, "2.4. API với NestJS và Fastify", 2, "ch2-4")
    for text in [
        "NestJS cung cấp dependency injection, module, controller, pipe, guard và interceptor. Fastify được chọn làm HTTP adapter để có pipeline gọn và khả năng xử lý cao. API dùng prefix /api, URI version v1 và DTO được kiểm tra bởi ValidationPipe với whitelist, forbidNonWhitelisted và transform; vì vậy trường không được định nghĩa bị từ chối thay vì âm thầm đi vào nghiệp vụ.",
        "Guard toàn cục xử lý JWT, vai trò và rate limit; decorator công khai cho phép các endpoint đọc như health hoặc tìm kiếm bỏ qua xác thực khi phù hợp. Swagger chỉ bật trong môi trường phát triển để hỗ trợ khám phá hợp đồng mà không mở bề mặt không cần thiết ở môi trường sản xuất. Request ID được gắn vào luồng xử lý để liên kết phản hồi lỗi với log.",
        "Thống kê tĩnh ghi nhận 30 controller và 186 endpoint, gồm 91 GET, 57 POST, 24 PATCH, 13 DELETE và 1 PUT. Phân bố này phù hợp hệ thống vừa có bề mặt đọc catalog lớn, vừa có quy trình tạo/duyệt nội dung và tác vụ quản trị.",
    ]:
        add_body(doc, text)
    add_listing(doc, "Hợp đồng phản hồi lỗi tham chiếu", [
        "HTTP/1.1 400 Bad Request",
        "{",
        '  "statusCode": 400,',
        '  "message": ["field must not be empty"],',
        '  "requestId": "7f6f4e4a-49b4-4c58-a487-3b8ec93260b0"',
        "}",
    ])

    add_heading(doc, "2.5. Prisma, PostgreSQL và tìm kiếm dữ liệu", 2, "ch2-5")
    for text in [
        "Prisma cung cấp lược đồ khai báo, client có kiểu và quy trình migration. Với 139 model, lợi ích quan trọng là hợp đồng dữ liệu được kiểm tra tập trung và quan hệ được biểu diễn rõ trong code. PostgreSQL 16 phù hợp dữ liệu quan hệ nhiều ràng buộc, đồng thời cho phép mở rộng tìm kiếm bằng pg_trgm, unaccent và lưu vector bằng pgvector.",
        "Tìm kiếm không chỉ là so khớp chuỗi. Truy vấn cần chuẩn hóa dấu, alias, tên model, hãng, chipset và từ khóa thuộc tính; kết quả phải phân trang, xếp hạng và quay lại thực thể đầy đủ. Kiến trúc cho phép dùng khả năng PostgreSQL làm nền và Meilisearch như bộ máy chuyên dụng tùy chọn. Dù đường tìm kiếm nào được chọn, kết quả cuối vẫn được hydrate từ catalog để duy trì một nguồn sự thật.",
        "Vector embedding hỗ trợ truy xuất ngữ nghĩa cho Wiki, thông số và câu hỏi AI. Vector không thay thế quan hệ hay citation: nó chỉ giúp chọn ngữ cảnh gần nghĩa. Thực thể gốc, phiên bản, model embedding và liên kết nguồn cần được lưu để có thể tái tạo hoặc vô hiệu hóa chỉ mục khi dữ liệu thay đổi.",
    ]:
        add_body(doc, text)

    add_heading(doc, "2.6. Redis, tác vụ nền và khả năng phục hồi", 2, "ch2-6")
    add_body(doc, "Redis được sử dụng cho phiên đăng nhập, cache và các trạng thái ngắn hạn cần truy cập nhanh. Refresh token gắn với session UUID lưu trong Redis giúp logout thực sự thu hồi phiên, thay vì chỉ xóa token ở trình duyệt. Cache phải có TTL và namespace theo phiên bản để tránh trả dữ liệu cũ sau khi catalog được xuất bản.")
    add_body(doc, "Worker xử lý cảnh báo giá, đồng bộ nguồn, lập chỉ mục hoặc gửi thông báo—những tác vụ không nên giữ request HTTP mở. Thiết kế job cần idempotent, có khóa chống chạy trùng, retry giới hạn và dead-letter/ghi lỗi rõ. Readiness kiểm tra DB và Redis giúp bộ điều phối chỉ gửi lưu lượng khi phụ thuộc cốt lõi sẵn sàng.")

    add_heading(doc, "2.7. Truy xuất tăng cường và AI có trích dẫn", 2, "ch2-7")
    for text in [
        "RAG kết hợp bước truy xuất ngữ cảnh với mô hình sinh nội dung. Câu hỏi được chuẩn hóa, tìm top-k đoạn hoặc thực thể phù hợp, sau đó API xây dựng prompt gồm chính sách trả lời, ngữ cảnh và định danh trích dẫn. Mô hình tạo câu trả lời; hệ thống kiểm tra định dạng và trả cả citation để giao diện hiển thị liên kết nguồn.",
        "Ba ranh giới kiểm soát chất lượng là retrieval, generation và presentation. Retrieval phải lọc theo trạng thái xuất bản và quyền; generation không được coi kiến thức ngoài context là dữ kiện chắc chắn; presentation cần phân biệt câu trả lời, cảnh báo thiếu bằng chứng và nguồn. Cache AI nên gắn với hash truy vấn, phiên bản dữ liệu và model để tránh dùng lại kết quả sau thay đổi quan trọng.",
        "Đánh giá AI không thể chỉ dùng test unit. Cần tập câu hỏi chuẩn, đáp án/nguồn mong đợi, thước đo citation precision, context recall, mức phù hợp và kiểm tra an toàn prompt injection. Phạm vi hiện tại đã có kiến trúc và endpoint, nhưng bộ benchmark định lượng toàn diện là hướng phát triển tiếp theo.",
    ]:
        add_body(doc, text)

    add_heading(doc, "2.8. Xác thực, phân quyền và bảo mật", 2, "ch2-8")
    add_body(doc, "Luồng xác thực phát hành access token ngắn hạn và refresh token gắn phiên. Mật khẩu được lưu dưới dạng băm; refresh cookie cần HttpOnly, Secure và SameSite phù hợp. RBAC phân biệt USER, EDITOR, MODERATOR và ADMIN; quyền chỉnh sửa catalog hoặc xuất bản Wiki được kiểm tra ở server, không dựa vào việc ẩn nút trên giao diện.")
    add_table_caption(doc, "2.2", "Biện pháp bảo mật theo lớp", "tab-2-2")
    add_table(doc,
              ["Lớp", "Rủi ro", "Biện pháp"],
              [
                  ("Trình duyệt", "XSS, đánh cắp token", "Escape nội dung, CSP/Helmet, refresh cookie HttpOnly, không lưu bí mật trong localStorage"),
                  ("HTTP/API", "Payload lạ, abuse", "HTTPS, CORS giới hạn, ValidationPipe, giới hạn kích thước và rate limit"),
                  ("Xác thực", "Credential stuffing, phiên tồn tại", "Hash mật khẩu, throttle, session Redis, rotate/thu hồi refresh"),
                  ("Phân quyền", "Leo thang đặc quyền", "JWT guard, Roles guard, deny-by-default và kiểm tra ownership"),
                  ("Dữ liệu", "Injection, mất toàn vẹn", "Prisma parameterization, ràng buộc DB, transaction, migration và backup"),
                  ("Webhook/API key", "Replay, lộ khóa", "Xác minh chữ ký, idempotency, lưu hash, scope và revoked_at"),
                  ("AI/RAG", "Prompt injection, lộ dữ liệu", "Lọc nguồn, tách system policy, không truy xuất tài liệu vượt quyền, citation bắt buộc"),
                  ("Vận hành", "Thiếu truy vết", "Request ID, audit log, metrics, cảnh báo và scrub dữ liệu nhạy cảm"),
              ], [1.05, 1.65, 3.5], 10.0)

    add_heading(doc, "2.9. PWA, quan sát hệ thống và triển khai", 2, "ch2-9")
    add_figure(doc, DIAGRAMS / "02-architecture.png", "2.2", "Kiến trúc logic nhiều lớp", "fig-2-2",
               "Sơ đồ kiến trúc logic mô tả lớp trình bày, API, nghiệp vụ, dữ liệu và hạ tầng.", page_break=True)
    add_body(doc, "PWA cải thiện khả năng cài đặt và phản hồi khi kết nối không ổn định, nhưng cache phải được thiết kế cẩn trọng để không làm stale dữ liệu giá hoặc trạng thái tài khoản. Quan sát hệ thống gồm health, liveness, readiness, metrics, structured log và request ID; audit log tập trung vào thay đổi có ảnh hưởng nghiệp vụ. Hai lớp này phục vụ các mục đích khác nhau và không nên trộn lẫn.")
    add_figure(doc, DIAGRAMS / "03-deployment.png", "2.3", "Mô hình triển khai tham chiếu", "fig-2-3",
               "Sơ đồ triển khai tham chiếu gồm thiết bị người dùng, dịch vụ web, API, worker, PostgreSQL, Redis và các dịch vụ ngoài.", page_break=True)
    add_table_caption(doc, "2.3", "Lựa chọn kiến trúc và đánh đổi", "tab-2-3")
    add_table(doc,
              ["Lựa chọn", "Lợi ích", "Đánh đổi / kiểm soát"],
              [
                  ("Monorepo", "Chia sẻ kiểu và pipeline", "Kho lớn; cần ranh giới package và cache build"),
                  ("Modular monolith API", "Giao dịch và triển khai đơn giản", "Cần kỷ luật phụ thuộc để tránh module chồng chéo"),
                  ("PostgreSQL trung tâm", "Toàn vẹn quan hệ và truy vấn linh hoạt", "Lược đồ lớn; cần index, migration và ownership rõ"),
                  ("Redis", "Phiên/cache nhanh", "Thêm phụ thuộc; phải thiết kế TTL, eviction và phục hồi"),
                  ("Meilisearch tùy chọn", "Tìm kiếm xếp hạng tốt", "Đồng bộ chỉ mục và chế độ fallback"),
                  ("RAG", "AI bám dữ liệu và có citation", "Chi phí, latency và yêu cầu benchmark chất lượng"),
                  ("Worker", "Tách tác vụ dài khỏi request", "Cần idempotency, retry, giám sát và xử lý lỗi"),
              ], [1.55, 2.1, 2.55], 10.1)

    add_heading(doc, "2.10. Kết luận chương", 2, "ch2-10")
    for text in [
        "Các công nghệ được lựa chọn tạo thành một kiến trúc TypeScript xuyên suốt, có dữ liệu quan hệ làm nền, bộ nhớ đệm và công cụ tìm kiếm là các năng lực bổ trợ, còn AI được đặt sau lớp truy xuất và trích dẫn. Điểm quan trọng không nằm ở danh sách framework mà ở cách các ranh giới và cơ chế kiểm soát được ghép thành những luồng có thể kiểm thử.",
        "Next.js và React đảm nhiệm lớp trình bày; NestJS/Fastify tổ chức hợp đồng HTTP và chính sách truy cập; Prisma/PostgreSQL bảo đảm tính toàn vẹn của dữ liệu nghiệp vụ; Redis hỗ trợ phiên, bộ nhớ đệm và điều phối ngắn hạn. Tiến trình nền tiếp nhận các công việc dài hoặc cần thử lại, qua đó giữ thời gian phản hồi của API trong giới hạn có thể kiểm soát. Cách phân vai này phù hợp với cấu trúc monorepo vì cho phép chia sẻ kiểu dữ liệu nhưng vẫn duy trì quyền sở hữu theo miền.",
        "Đối với tìm kiếm và AI, thiết kế ưu tiên dữ liệu đã xuất bản, nguồn có thể kiểm chứng và cơ chế suy giảm có kiểm soát. Kết quả sinh không được xem là dữ liệu gốc; hệ thống cần lưu ngữ cảnh, phiên bản chỉ mục, trích dẫn và dấu vết đánh giá. Đối với bảo mật, xác thực chỉ là lớp đầu; quyết định truy cập còn phải dựa trên vai trò, phạm vi tài nguyên, trạng thái nội dung, giới hạn tần suất và nhật ký kiểm toán.",
        "Các lựa chọn trên tạo ra cơ sở kỹ thuật để chuyển yêu cầu nghiệp vụ thành mô hình thiết kế có thể triển khai. Chương 3 tiếp tục cụ thể hóa bằng tác nhân, ca sử dụng, luồng tuần tự, quy trình nền, hợp đồng API và các sơ đồ thực thể–liên kết theo miền dữ liệu.",
    ]:
        add_body(doc, text)


def use_case_table(doc, number, key, title, data):
    doc.add_page_break()
    add_table_caption(doc, number, f"Đặc tả {data['id']} – {title}", key)
    rows = [
        ("Mã use case", data["id"]),
        ("Mục tiêu", data["goal"]),
        ("Tác nhân", data["actors"]),
        ("Tiền điều kiện", data["pre"]),
        ("Kích hoạt", data["trigger"]),
        ("Luồng chính", data["main"]),
        ("Luồng thay thế / ngoại lệ", data["alt"]),
        ("Hậu điều kiện", data["post"]),
        ("Dữ liệu và kiểm soát", data["controls"]),
    ]
    add_table(doc, ["Thuộc tính", "Nội dung"], rows, [1.45, 4.75], 10.5)


def add_chapter_3(doc):
    add_heading(doc, "CHƯƠNG 3. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", 1, "ch3")
    add_heading(doc, "3.1. Tác nhân và mô hình use case", 2, "ch3-1")
    add_body(doc, "Phân tích tác nhân dựa trên quyền và mục tiêu nghiệp vụ, không đồng nhất tác nhân với một màn hình. Khách truy cập có thể thực hiện luồng đọc; người dùng kế thừa quyền đọc và có dữ liệu cá nhân; ba vai trò nội bộ có đặc quyền tăng dần; đối tác B2B tương tác bằng API key thay vì phiên trình duyệt.")
    add_table_caption(doc, "3.1", "Tác nhân hệ thống", "tab-3-1")
    add_table(doc,
              ["Tác nhân", "Xác thực", "Khả năng chính", "Giới hạn"],
              [
                  ("Khách truy cập", "Không bắt buộc", "Tìm kiếm, lọc, xem chi tiết, so sánh, khuyến nghị công khai", "Không lưu dữ liệu cá nhân lâu dài"),
                  ("Người dùng", "JWT + refresh session", "Wishlist, cảnh báo, thông báo, đóng góp Wiki", "Chỉ thao tác dữ liệu của mình; nội dung chờ duyệt"),
                  ("Biên tập viên", "JWT, role EDITOR", "Soạn catalog, quản lý nguồn, đối soát dữ liệu", "Không quản trị người dùng/hệ thống ngoài phạm vi"),
                  ("Kiểm duyệt viên", "JWT, role MODERATOR", "Duyệt thay đổi và phiên bản Wiki", "Quyết định được ghi audit"),
                  ("Quản trị viên", "JWT, role ADMIN", "Quản trị quyền, gói, API key, webhook và vận hành", "Thao tác nhạy cảm phải có log và xác nhận"),
                  ("Đối tác B2B", "API key + scope", "Đọc dữ liệu theo hợp đồng và hạn mức", "Không dùng endpoint nội bộ; có quota và thu hồi"),
                  ("Worker", "Danh tính dịch vụ", "Đồng bộ, cảnh báo, thông báo, lập chỉ mục", "Không nhận tương tác trực tiếp; job phải idempotent"),
              ], [1.1, 1.4, 2.45, 1.25], 9.5)
    add_figure(doc, DIAGRAMS / "04-usecase-user.png", "3.1", "Use case phía người dùng", "fig-3-1",
               "Sơ đồ use case phía người dùng mô tả tìm kiếm, xem chi tiết, so sánh, khuyến nghị, AI, Wiki, wishlist và cảnh báo.", page_break=True)
    add_figure(doc, DIAGRAMS / "05-usecase-admin.png", "3.2", "Use case biên tập và quản trị", "fig-3-2",
               "Sơ đồ use case nội bộ mô tả biên tập catalog, duyệt, quản lý người dùng, audit, gói thuê bao và theo dõi hệ thống.", page_break=True)
    add_table_caption(doc, "3.2", "Danh mục use case", "tab-3-2")
    add_table(doc,
              ["Mã", "Tên", "Tác nhân chính", "Kết quả"],
              [
                  ("UC-01", "Tìm kiếm và lọc thiết bị", "Khách/Người dùng", "Danh sách xếp hạng, phân trang và có bộ lọc"),
                  ("UC-02", "So sánh thiết bị", "Khách/Người dùng", "Ma trận thuộc tính chuẩn hóa"),
                  ("UC-03", "Khuyến nghị và hỏi AI", "Khách/Người dùng", "Gợi ý/câu trả lời có lý do và citation"),
                  ("UC-04", "Đăng nhập và quản lý phiên", "Người dùng nội bộ", "Access token và refresh session có thể thu hồi"),
                  ("UC-05", "Wishlist và cảnh báo giá", "Người dùng", "Theo dõi variant và notification khi đạt ngưỡng"),
                  ("UC-06", "Đóng góp và duyệt Wiki", "Người dùng/Moderator", "Revision được xuất bản hoặc từ chối có lý do"),
                  ("UC-07", "Thu thập và duyệt catalog", "Editor/Moderator/Worker", "Dữ liệu chuẩn hóa có nguồn được xuất bản"),
                  ("UC-08", "Khai thác API B2B", "Đối tác/Admin", "Truy cập theo scope, quota và usage"),
              ], [0.75, 2.2, 1.55, 1.7], 9.8)

    add_heading(doc, "3.2. Đặc tả use case", 2, "ch3-2")
    add_body(doc, "Tám use case dưới đây đại diện cho hành trình đọc, cá nhân hóa, cộng tác, quản trị dữ liệu và tích hợp. Luồng được viết ở mức nghiệp vụ; đường dẫn API cụ thể có thể thay đổi nhưng tiền/hậu điều kiện và kiểm soát phải được giữ.")
    use_case_table(doc, "3.3", "tab-3-3", "Tìm kiếm và lọc thiết bị", {
        "id": "UC-01", "goal": "Tìm nhanh thiết bị hoặc phần cứng phù hợp theo từ khóa, loại, hãng và thuộc tính.",
        "actors": "Khách truy cập; Người dùng.",
        "pre": "Catalog có dữ liệu đã xuất bản; dịch vụ tìm kiếm ở chế độ PostgreSQL hoặc Meilisearch sẵn sàng.",
        "trigger": "Tác nhân nhập từ khóa, chọn bộ lọc hoặc mở URL có query parameter.",
        "main": "1) Web chuẩn hóa input và gửi GET search.\n2) API kiểm tra tham số, giới hạn trang và sort.\n3) Search tìm ID xếp hạng theo text/alias.\n4) Catalog hydrate model, variant, media và facet.\n5) API trả DTO phân trang.\n6) Web hiển thị kết quả, bộ lọc và trạng thái URL.",
        "alt": "A1) Từ khóa rỗng: trả danh mục theo filter.\nA2) Không có kết quả: gợi ý bỏ bớt filter hoặc alias gần đúng.\nE1) Search chuyên dụng lỗi: fallback sang PostgreSQL nếu cấu hình cho phép.\nE2) Tham số sai: trả 400 cùng request ID.",
        "post": "Không thay đổi catalog; có thể ghi search log và latency mà không lưu dữ liệu nhạy cảm.",
        "controls": "Whitelist filter; giới hạn page size; chỉ trả bản ghi published; query được parameter hóa; cache theo query chuẩn hóa.",
    })
    use_case_table(doc, "3.4", "tab-3-4", "So sánh thiết bị", {
        "id": "UC-02", "goal": "Đặt từ hai đến một số hữu hạn thiết bị cạnh nhau trên cùng hệ thuộc tính.",
        "actors": "Khách truy cập; Người dùng.",
        "pre": "Các slug hợp lệ và model có ít nhất một variant được xuất bản.",
        "trigger": "Tác nhân chọn nút So sánh từ danh sách/chi tiết hoặc mở URL compare.",
        "main": "1) Web duy trì danh sách slug không trùng.\n2) API nhận danh sách và kiểm tra số lượng.\n3) Service tải model, variant và phần cứng liên quan.\n4) Thuộc tính được ánh xạ vào schema so sánh theo category.\n5) Giá trị được chuẩn hóa đơn vị và đánh dấu thiếu dữ liệu.\n6) Web hiển thị ma trận, điểm khác biệt và nguồn khi có.",
        "alt": "A1) Chỉ có một thiết bị: yêu cầu chọn thêm.\nA2) Khác category không tương thích: dùng tập thuộc tính chung hoặc cảnh báo.\nE1) Slug không tồn tại: bỏ mục lỗi và thông báo.\nE2) Dữ liệu thiếu: hiển thị 'Chưa có dữ liệu', không suy diễn.",
        "post": "URL chứa lựa chọn có thể chia sẻ; không thay đổi dữ liệu người dùng nếu chưa lưu.",
        "controls": "Giới hạn số model; thứ tự ổn định; chuẩn hóa đơn vị; citation gắn thuộc tính; tránh N+1 bằng include/batch.",
    })
    use_case_table(doc, "3.5", "tab-3-5", "Nhận khuyến nghị và hỏi AI", {
        "id": "UC-03", "goal": "Chuyển nhu cầu tự nhiên hoặc bộ tiêu chí thành gợi ý có thể giải thích và câu trả lời có nguồn.",
        "actors": "Khách truy cập; Người dùng; AI service.",
        "pre": "Catalog/Wiki đã lập chỉ mục; chính sách prompt và cấu hình model hợp lệ.",
        "trigger": "Tác nhân hoàn thành biểu mẫu khuyến nghị hoặc gửi câu hỏi.",
        "main": "1) API validate nhu cầu/câu hỏi.\n2) Bộ truy xuất tạo query và lấy top-k context theo quyền.\n3) Catalog tải dữ kiện/citation tương ứng.\n4) AI service tạo prompt ràng buộc.\n5) LLM trả nội dung có tham chiếu ID.\n6) API kiểm tra, lưu log/cache phù hợp và trả câu trả lời cùng citation.\n7) Web hiển thị kết quả, lý do và liên kết nguồn.",
        "alt": "A1) Không đủ context: trả lời giới hạn và đề nghị làm rõ.\nA2) Model ngoài danh mục: nêu rõ không tìm thấy.\nE1) Dịch vụ LLM lỗi: trả thông báo có thể thử lại, không giả câu trả lời.\nE2) Citation không hợp lệ: loại phần không kiểm chứng hoặc đánh dấu.",
        "post": "Có bản ghi đo chất lượng/latency nếu chính sách cho phép; catalog không bị sửa bởi câu trả lời.",
        "controls": "Giới hạn độ dài; lọc prompt injection; chỉ context published; timeout; cache theo phiên bản; citation bắt buộc cho dữ kiện.",
    })
    use_case_table(doc, "3.6", "tab-3-6", "Đăng nhập và quản lý phiên", {
        "id": "UC-04", "goal": "Xác thực an toàn và cho phép thu hồi phiên độc lập với access token.",
        "actors": "Người dùng; Biên tập viên; Moderator; Admin.",
        "pre": "Tài khoản hoạt động; PostgreSQL và Redis sẵn sàng.",
        "trigger": "Tác nhân gửi email/mật khẩu hoặc access token hết hạn cần làm mới.",
        "main": "1) API validate credential.\n2) Tìm user và so mật khẩu băm.\n3) Tạo session UUID trong Redis với TTL.\n4) Phát access token ngắn hạn và refresh token.\n5) Web nhận hồ sơ và lưu trạng thái.\n6) Khi refresh, API xác minh chữ ký và session.\n7) Logout xóa/thu hồi session và cookie.",
        "alt": "E1) Sai credential: phản hồi chung, không tiết lộ email tồn tại.\nE2) Session hết hạn/thu hồi: yêu cầu đăng nhập lại.\nE3) Tài khoản khóa: từ chối và audit.\nE4) Redis không sẵn sàng: readiness fail; không cấp phiên không thể quản lý.",
        "post": "Phiên hợp lệ được lưu có TTL hoặc đã bị thu hồi; sự kiện đăng nhập đặc quyền có thể được audit.",
        "controls": "Rate limit; hash mật khẩu; cookie HttpOnly/Secure; rotate token; RBAC server-side; scrub token khỏi log.",
    })
    use_case_table(doc, "3.7", "tab-3-7", "Wishlist và cảnh báo giá", {
        "id": "UC-05", "goal": "Lưu variant quan tâm và nhận thông báo khi giá đạt điều kiện.",
        "actors": "Người dùng; Worker; Notification service.",
        "pre": "Người dùng đăng nhập; variant tồn tại; kênh thông báo được cấu hình.",
        "trigger": "Người dùng thêm wishlist hoặc đặt target price.",
        "main": "1) API kiểm tra ownership, variant, tiền tệ và ngưỡng.\n2) Lưu wishlist/alert duy nhất.\n3) Worker định kỳ đọc giá mới.\n4) So khớp điều kiện và chống gửi lặp.\n5) Tạo notification và delivery.\n6) Gửi qua kênh được bật.\n7) Ghi trạng thái delivery và lần kích hoạt.",
        "alt": "A1) Giá đã dưới ngưỡng tại lúc tạo: thông báo ngay theo chính sách.\nA2) Người dùng tắt alert: worker bỏ qua.\nE1) Kênh gửi lỗi: retry giới hạn, ghi failed.\nE2) Giá thiếu/stale: không kích hoạt, ghi lý do.",
        "post": "Alert có trạng thái và last_triggered_at; delivery có lịch sử để chống lặp và hỗ trợ.",
        "controls": "Unique user+variant; decimal/currency chuẩn; idempotency key; quiet hours; preference và unsubscribe.",
    })
    use_case_table(doc, "3.8", "tab-3-8", "Đóng góp và duyệt Wiki", {
        "id": "UC-06", "goal": "Cho phép cộng tác nội dung mà vẫn duy trì lịch sử, nguồn và quyền xuất bản.",
        "actors": "Người dùng; Biên tập viên; Kiểm duyệt viên.",
        "pre": "Tác nhân đăng nhập; bài viết hoặc quyền tạo mới hợp lệ.",
        "trigger": "Tác nhân tạo/sửa bài Wiki và gửi duyệt.",
        "main": "1) Web tải current revision và citation.\n2) Tác nhân soạn nội dung, thêm nguồn.\n3) API validate và tạo revision draft.\n4) Gửi revision vào hàng đợi review.\n5) Moderator xem diff, nguồn và bình luận.\n6) Phê duyệt tạo current revision mới hoặc từ chối có lý do.\n7) Sự kiện xuất bản làm mới search/embedding.",
        "alt": "A1) Xung đột revision: yêu cầu merge trên phiên bản mới.\nA2) Yêu cầu bổ sung nguồn: trả về draft.\nE1) Citation không hợp lệ: không cho xuất bản.\nE2) Tác nhân thiếu role: trả 403 và audit.",
        "post": "Mọi revision được giữ; bài published trỏ current_revision; quyết định review có actor và thời gian.",
        "controls": "Optimistic concurrency; sanitize rich text; RBAC; audit; không sửa trực tiếp revision đã xuất bản.",
    })
    use_case_table(doc, "3.9", "tab-3-9", "Thu thập và duyệt catalog", {
        "id": "UC-07", "goal": "Đưa dữ liệu nguồn vào catalog qua pipeline có kiểm soát và truy nguyên.",
        "actors": "Worker; Biên tập viên; Kiểm duyệt viên.",
        "pre": "Data source được đăng ký, có chính sách thu thập và parser phù hợp.",
        "trigger": "Lịch job, webhook hoặc thao tác đồng bộ thủ công.",
        "main": "1) Worker fetch và lưu raw page/content hash.\n2) Parser trích xuất candidate.\n3) Chuẩn hóa alias, đơn vị và quan hệ.\n4) So với phiên bản catalog, tạo diff kèm citation.\n5) Editor đối soát và chỉnh candidate.\n6) Moderator phê duyệt/từ chối.\n7) Transaction xuất bản, audit và phát sự kiện reindex.",
        "alt": "A1) Content hash không đổi: kết thúc không tạo diff.\nA2) Nguồn xung đột: giữ candidate và yêu cầu nguồn thứ hai.\nE1) Parser lỗi: lưu raw page/log, không tác động catalog.\nE2) Transaction lỗi: rollback toàn bộ.",
        "post": "Catalog chỉ thay đổi khi được duyệt; nguồn, raw page, diff và quyết định được lưu.",
        "controls": "Rate/robots policy; hash; schema validation; transaction; idempotency; audit và quyền theo vai trò.",
    })
    use_case_table(doc, "3.10", "tab-3-10", "Khai thác API B2B", {
        "id": "UC-08", "goal": "Cho đối tác truy cập dữ liệu theo scope, gói và hạn mức có thể theo dõi.",
        "actors": "Đối tác B2B; Quản trị viên.",
        "pre": "Đối tác có subscription/API key hoạt động và scope phù hợp.",
        "trigger": "Đối tác gửi request có API key tới endpoint cho phép.",
        "main": "1) Gateway/guard băm key và tìm bản ghi.\n2) Kiểm tra revoked_at, subscription, scope và quota.\n3) Validate tham số và gọi service catalog.\n4) Trả DTO versioned.\n5) Ghi usage theo đơn vị và route.\n6) Trả header quota/request ID.\n7) Admin xem usage hoặc thu hồi/rotate key.",
        "alt": "E1) Key sai/thu hồi: 401.\nE2) Thiếu scope: 403.\nE3) Vượt quota: 429 và retry guidance.\nE4) Gói hết hạn: chặn và phát sự kiện billing nếu cần.",
        "post": "Usage được cộng idempotent; không lưu key dạng rõ; đối tác có thể quan sát hạn mức.",
        "controls": "Key chỉ hiển thị một lần; lưu hash; scope tối thiểu; rate limit; audit rotate/revoke; API versioning.",
    })

    add_heading(doc, "3.3. Thiết kế sequence diagram", 2, "ch3-3")
    add_body(doc, "Sequence diagram tập trung vào trust boundary và nơi dữ liệu được validate. Login flow cho thấy refresh token không tự đủ để cấp session mới; Redis session là điều kiện bắt buộc. Search flow tách ranking ID khỏi hydrate entity. AI flow tách retrieval, catalog và generation để giữ citation.")
    add_figure(doc, DIAGRAMS / "06-sequence-auth.png", "3.3", "Sequence diagram đăng nhập và refresh session", "fig-3-3",
               "Sequence diagram mô tả user, web, Auth API, PostgreSQL và Redis trong login và refresh flow.")
    add_figure(doc, DIAGRAMS / "08-sequence-search-compare.png", "3.4", "Sequence diagram tìm kiếm và so sánh", "fig-3-4",
               "Sequence diagram mô tả search, hydrate catalog và trả comparison matrix.", page_break=True)
    add_figure(doc, DIAGRAMS / "07-sequence-ai-rag.png", "3.5", "Sequence diagram trợ lý AI theo RAG", "fig-3-5",
               "Sequence diagram mô tả vector retrieval, catalog lookup, prompt construction, LLM call và citation response.", page_break=True)
    add_body(doc, "Điểm chung của ba luồng là Web không truy cập dữ liệu nền trực tiếp. API chịu trách nhiệm validate, authorize và tạo DTO; các dịch vụ hạ tầng trả dữ liệu kỹ thuật, còn quyết định nghiệp vụ nằm trong module. Nhờ đó, client không cần biết cấu trúc 139 bảng và không thể bỏ qua quy tắc bằng cách gọi DB/search trực tiếp.")

    add_heading(doc, "3.4. Thiết kế background workflow", 2, "ch3-4")
    add_body(doc, "Hai background workflow có state và yêu cầu idempotency. Price alert flow phải phân biệt current price với threshold event, tôn trọng user preference và lưu delivery attempt. Ingestion pipeline phải giữ raw page ngay cả khi parser thất bại để hỗ trợ debugging và reprocessing.")
    add_figure(doc, DIAGRAMS / "09-flow-price-alert.png", "3.6", "Price alert workflow", "fig-3-6",
               "Workflow mô tả tạo alert, worker đọc price point, match threshold, tạo notification và ghi delivery.", page_break=True)
    add_figure(doc, DIAGRAMS / "10-flow-ingestion.png", "3.7", "Ingestion workflow", "fig-3-7",
               "Workflow mô tả data source, raw page, extraction, normalization, review, publish và indexing.", page_break=True)

    add_heading(doc, "3.5. Thiết kế dữ liệu", 2, "ch3-5")
    for text in [
        "Lược đồ Prisma được phân tích có 139 model. Một ERD duy nhất với toàn bộ bảng sẽ không đọc được trên khổ A4, vì vậy báo cáo chia thành bốn miền và chỉ hiển thị thực thể/thuộc tính tiêu biểu. Danh sách vật lý đầy đủ được đưa vào Phụ lục B. Các bảng dùng snake_case và quan hệ được định nghĩa rõ trong Prisma.",
        "Nguyên tắc thiết kế là giữ thuộc tính cốt lõi ở cột có kiểu, dùng bảng nối cho quan hệ N–N, tách dữ liệu lịch sử/phiên bản khỏi bản hiện hành, và tách citation khỏi thực thể để một nguồn có thể chứng minh nhiều mệnh đề. JSON chỉ phù hợp metadata linh hoạt, không thay thế khóa ngoại cho quan hệ cần truy vấn hoặc ràng buộc.",
    ]:
        add_body(doc, text)
    add_table_caption(doc, "3.11", "Phân nhóm 139 mô hình dữ liệu", "tab-3-11")
    add_table(doc,
              ["Miền", "Nhóm thực thể", "Mục đích"],
              [
                  ("Core catalog", "organizations, device_categories, product_families, device_models, device_variants", "Định danh sản phẩm và variant chuẩn hóa"),
                  ("Hardware/scoring", "chipsets, display_units, join tables, scoring_profiles, scorecards, benchmarks", "Mô tả component, cấu hình và kết quả đo"),
                  ("Content/crawler/AI", "sources, citations, Wiki, revision, raw_pages, embeddings, AI query cache, search logs", "Traceability, cộng tác và semantic retrieval"),
                  ("User engagement/commerce", "users, wishlists, alerts, notifications, subscriptions, API keys, affiliate, webhooks", "Cá nhân hóa, doanh thu và vận hành"),
              ], [1.25, 3.25, 1.7], 9.6)
    add_figure(doc, DIAGRAMS / "11-erd-catalog.png", "3.8", "ERD nhóm catalog", "fig-3-8",
               "ERD thể hiện các physical foreign key tiêu biểu giữa organizations, device_categories, product_families, device_models, device_variants, release_statuses, currencies và variant_price_history.", page_break=True)
    add_body(doc, "Catalog tách device_model khỏi device_variant vì tên sản phẩm và marketing content có thể dùng chung, trong khi SKU, màu, launch price và currency thuộc variant. device_model_aliases hỗ trợ search theo tên khác mà không nhân bản model; variant_price_history lưu price point theo thời gian và tham chiếu trực tiếp device_variant.")
    add_figure(doc, DIAGRAMS / "12-erd-hardware.png", "3.9", "ERD nhóm hardware và scoring", "fig-3-9",
               "ERD thể hiện join table variant_chipsets, variant_displays cùng scorecard và benchmark của device_variant.", page_break=True)
    add_body(doc, "Hardware component có thể tái sử dụng giữa nhiều device_variant. Các join table giữ role hoặc position của lần gắn component. scoring_profiles tách scoring policy khỏi variant_scorecards; device_variant_benchmarks tách benchmark definition khỏi measured result để giữ traceability theo từng lần đo.")
    add_figure(doc, DIAGRAMS / "13-erd-content-ai.png", "3.10", "ERD nhóm content, crawler và AI/search", "fig-3-10",
               "ERD thể hiện sources/citations, Wiki article/revision, crawler raw_pages và logical reference tới embeddings.", page_break=True)
    add_body(doc, "wiki_revisions là immutable revision; wiki_articles giữ current_revision_id. wiki_article_citations nối article với citation. data_sources và raw_pages thuộc crawler domain. embeddings sử dụng polymorphic logical reference entity_type/entity_id; connector nét đứt nhấn mạnh đây không phải physical foreign key trong Prisma schema.")
    add_figure(doc, DIAGRAMS / "14-erd-user-commerce.png", "3.11", "ERD nhóm user engagement, subscription và B2B", "fig-3-11",
               "ERD thể hiện user ownership đối với wishlist, price alert, notification, subscription và API key.", page_break=True)
    add_body(doc, "notifications biểu diễn notification intent; notification_deliveries biểu diễn từng delivery attempt theo channel. Cách tách này cho phép retry và đo delivery status mà không tạo notification trùng. subscriptions tham chiếu subscription_plans; api_keys chỉ lưu key_hash và scopes; price_alerts tham chiếu đồng thời users và device_variants.")

    add_heading(doc, "3.6. Thiết kế API và hợp đồng dữ liệu", 2, "ch3-6")
    add_body(doc, "API sử dụng URI version v1 để cho phép thay đổi hợp đồng có kiểm soát. Resource path ưu tiên danh từ; POST dùng cho tạo hoặc action có side effect; PATCH cho cập nhật một phần; DELETE cho thu hồi/xóa theo chính sách. DTO không lộ trực tiếp cấu trúc Prisma và chỉ trả trường cần thiết theo quyền.")
    add_table_caption(doc, "3.12", "Quy ước hợp đồng API", "tab-3-12")
    add_table(doc,
              ["Khía cạnh", "Quy ước", "Lý do"],
              [
                  ("Đường dẫn cơ sở", "/api/v1", "Tách hạ tầng và phiên bản URI"),
                  ("Dữ liệu đầu vào", "DTO + ValidationPipe; whitelist/forbid/transform", "Không cho trường lạ đi vào lớp dịch vụ"),
                  ("Xác thực", "JWT/phiên hoặc khóa API; @Public khi cần", "Mặc định yêu cầu xác thực, ngoại lệ được khai báo rõ"),
                  ("Phân quyền", "ADMIN / EDITOR / MODERATOR", "Thực hiện kiểm soát tại máy chủ"),
                  ("Phân trang", "page/pageSize hoặc cursor có giới hạn", "Tránh tải tập dữ liệu không giới hạn"),
                  ("Lỗi", "Mã trạng thái HTTP + thông báo + mã định danh yêu cầu", "Máy khách xử lý nhất quán và hỗ trợ truy vết"),
                  ("Tính lũy đẳng", "Khóa hoặc trạng thái duy nhất cho webhook/tác vụ/thao tác", "Chống thực hiện lặp khi thử lại"),
                  ("Quan sát", "Độ trễ, trạng thái, mẫu tuyến; không ghi bí mật vào nhật ký", "Đo vận hành mà vẫn bảo vệ dữ liệu"),
              ], [1.15, 3.0, 2.05], 10.0)

    add_heading(doc, "3.7. Thiết kế an toàn và kiểm soát chất lượng", 2, "ch3-7")
    add_body(doc, "Mô hình đe dọa tập trung vào bốn nhóm tài sản: tài khoản và phiên đăng nhập; dữ liệu danh mục có nguồn; khóa API và dữ liệu sử dụng; ngữ cảnh và câu trả lời AI. Biên tin cậy xuất hiện giữa trình duyệt với API, API với dịch vụ ngoài, tiến trình nền với nguồn dữ liệu và webhook với hệ thống thanh toán. Tại mỗi biên cần xác thực nguồn, giới hạn dữ liệu đầu vào, quy định thời gian chờ và ghi nhật ký sự kiện.")
    add_bullets(doc, [
        "Quyền sở hữu: người dùng chỉ được sửa danh sách yêu thích và cảnh báo của mình; biên tập viên không mặc nhiên có quyền quản trị tài khoản.",
        "Tính toàn vẹn: sử dụng giao dịch cơ sở dữ liệu khi xuất bản nhiều bảng; phiên bản đã xuất bản là bất biến; webhook bảo đảm tính lũy đẳng.",
        "Tính bí mật: mật khẩu, khóa API và mã xác thực chỉ được lưu dưới dạng băm hoặc bí mật; nhật ký phải loại bỏ dữ liệu nhạy cảm.",
        "Tính khả dụng: giới hạn tần suất, kích thước trang và tải tin; tiến trình nền thử lại theo khoảng lùi và kiểm tra mức sẵn sàng của phụ thuộc.",
        "Chất lượng dữ liệu: lưu mã băm nội dung, trích dẫn nguồn, trạng thái kiểm duyệt, nhật ký hoạt động và thống kê trường còn thiếu.",
        "Chất lượng AI: chỉ dùng ngữ cảnh đã xuất bản, kiểm tra trích dẫn nguồn, quy định thời gian chờ, có phương án dự phòng và bộ câu hỏi đánh giá chuẩn.",
    ])
    add_body(doc, "Kiểm soát chất lượng được đặt trước và sau hiện thực: lint/typecheck/schema validation ngăn lỗi cấu trúc; unit test kiểm tra quy tắc; readiness kiểm tra kết nối; ảnh giao diện xác nhận luồng; còn đo tải, E2E và kiểm thử bảo mật động là các lớp cần bổ sung trước sản xuất.")

    add_heading(doc, "3.8. Kết luận chương", 2, "ch3-8")
    add_body(doc, "Chương 3 đã chuyển các yêu cầu thành tám ca sử dụng, ba luồng tuần tự, hai quy trình nền và bốn sơ đồ thực thể–liên kết theo miền. Thiết kế nhấn mạnh ranh giới quyền, dữ liệu đã xuất bản, tính lũy đẳng, trích dẫn nguồn và khả năng truy vết. Chương 4 trình bày quá trình hiện thực các thành phần này trong mã nguồn, giao diện và kết quả đánh giá.")


def add_chapter_4(doc, evidence):
    add_heading(doc, "CHƯƠNG 4. PHÁT TRIỂN VÀ KẾT QUẢ", 1, "ch4")
    add_heading(doc, "4.1. Môi trường và quy trình phát triển", 2, "ch4-1")
    for text in [
        "Dự án yêu cầu Node.js từ phiên bản 22.11, pnpm 9.15 cùng PostgreSQL và Redis. Quy trình phát triển gồm cài đặt các thư viện phụ thuộc tại thư mục gốc, cấu hình biến môi trường từ tệp mẫu, khởi động hạ tầng, thực hiện di trú và tạo dữ liệu mẫu khi cần, sau đó chạy giao diện web, API và tiến trình nền bằng Turborepo. Mỗi ứng dụng vẫn có lệnh riêng để hỗ trợ gỡ lỗi độc lập.",
        "Mã nguồn TypeScript được kiểm tra quy tắc viết mã và kiểu dữ liệu; bộ kiểm thử đơn vị của API được thực thi bằng Jest; công cụ dòng lệnh Prisma kiểm tra lược đồ dữ liệu. Swagger hỗ trợ tra cứu API trong môi trường phát triển. Các điểm kiểm tra trạng thái tổng hợp, trạng thái tiến trình, mức sẵn sàng và chỉ số giám sát được tách riêng để phân biệt tiến trình đang chạy với khả năng tiếp nhận lưu lượng.",
    ]:
        add_body(doc, text)
    add_table_caption(doc, "4.1", "Môi trường phần mềm", "tab-4-1")
    add_table(doc,
              ["Thành phần", "Yêu cầu / cấu hình", "Mục đích đánh giá"],
              [
                  ("Node.js", ">= 22.11.0", "Môi trường thực thi thống nhất cho Next.js, NestJS và chuỗi công cụ"),
                  ("pnpm", "9.15.0", "Cài đặt kho mã nguồn và tệp khóa có thể tái lập"),
                  ("PostgreSQL", "16 + pgvector/pg_trgm/unaccent", "Lược đồ quan hệ, tìm kiếm và véc-tơ nhúng"),
                  ("Redis", "Kết nối bằng biến môi trường", "Phiên, bộ nhớ đệm và điều phối"),
                  ("Web", "127.0.0.1:3000 khi kiểm tra cục bộ", "22 tuyến giao diện và PWA"),
                  ("API", "127.0.0.1:4000, cơ sở /api/v1", "186 điểm cuối, kiểm tra trạng thái và Swagger trong môi trường phát triển"),
                  ("Kiểm thử", "Jest trong phạm vi API", "36 bộ / 189 ca kiểm thử trong lần đánh giá"),
                  ("Prisma", "kiểm tra/tạo mã/di trú", "Kiểm tra lược đồ và trình khách Prisma"),
              ], [1.2, 2.45, 2.55], 10.2)

    add_heading(doc, "4.2. Hiện thực các mô-đun chính", 2, "ch4-2")
    add_body(doc, "Kết quả khảo sát mã nguồn cho thấy hệ thống không dồn toàn bộ chức năng vào một API tổng quát. Các mô-đun được tổ chức theo miền nghiệp vụ và sử dụng lớp dịch vụ để phối hợp Prisma, Redis, chức năng tìm kiếm hoặc dịch vụ bên ngoài. Bảng 4.2 tóm tắt các chức năng đã được cài đặt; Phụ lục A liệt kê toàn bộ bộ điều khiển API.")
    add_table_caption(doc, "4.2", "Mô-đun hiện thực chính", "tab-4-2")
    add_table(doc,
              ["Mô-đun", "Chức năng", "Kỹ thuật nổi bật"],
              [
                  ("Catalog/Hardware", "Model, variant, linh kiện, media, alias, score", "Prisma relation, slug, filter, DTO"),
                  ("Search/Compare", "Tìm text/facet và dựng ma trận so sánh", "pg_trgm/unaccent; Meilisearch tùy chọn; batch hydrate"),
                  ("Auth/User", "Register/login/refresh/me/logout và hồ sơ", "JWT, password hash, Redis session, RBAC"),
                  ("AI", "Ask, chat, recommendation, context/citation", "Retrieval, embedding, cache, model adapter"),
                  ("Wiki/Moderation", "Article, revision, citation, comment và review", "Versioning, diff, trạng thái và role"),
                  ("Wishlist/Alerts", "Theo dõi variant, ngưỡng giá và notification", "Ownership, worker, idempotency, delivery"),
                  ("Affiliate/Price", "Đối tác, link, click, lịch sử và insight", "Sync, inspect, audit và attribution"),
                  ("Billing/B2B", "Subscription, webhook, API key và usage", "Stripe, signature, hash key, scope/quota"),
                  ("Admin/Observability", "Dashboard, audit, health và metrics", "Guard, request ID, readiness và log"),
              ], [1.35, 2.65, 2.2], 9.8)
    add_body(doc, "Một nguyên tắc hiện thực đáng chú ý là public endpoint và role không được suy ra từ tên đường dẫn. Decorator/guard trên server xác định quyền; đối với các thao tác chỉnh sửa hoặc quản trị, service vẫn phải kiểm tra ownership/trạng thái nghiệp vụ. Điều này tránh tình huống một endpoint có URL công khai nhưng vô tình cho phép hành động đặc quyền.")

    add_heading(doc, "4.3. Kết quả giao diện", 2, "ch4-3")
    add_body(doc, "Giao diện được đánh giá ở độ phân giải 1440 × 900 trên hệ thống vận hành cục bộ với dữ liệu hiện có. Các hình dưới đây được chụp trực tiếp từ ứng dụng và minh họa những chức năng chính. Tại thời điểm ngày 05/08/2026, trang danh mục hiển thị 282 mẫu thiết bị, 39 hãng và 8 loại thiết bị; các số liệu này phản ánh trạng thái dữ liệu tại thời điểm đánh giá.")
    add_table_caption(doc, "4.3", "Bản đồ route giao diện", "tab-4-3")
    route_items = [(route, {
        "/": "Trang chủ/điều hướng khám phá", "/admin": "Catalog Studio và quản trị", "/ai": "Trợ lý AI",
        "/alerts": "Cảnh báo giá", "/api-access": "API key và usage", "/billing": "Gói thuê bao",
        "/compare": "So sánh", "/dashboard": "Bảng điều khiển cá nhân", "/devices": "Danh mục thiết bị",
        "/devices/[slug]": "Chi tiết model", "/hardware/[kind]/[slug]": "Chi tiết phần cứng", "/login": "Đăng nhập",
        "/notifications": "Thông báo", "/offline": "Trạng thái PWA offline", "/recommend": "Khuyến nghị",
        "/register": "Đăng ký", "/search": "Tìm kiếm", "/wiki": "Danh mục Wiki",
        "/wiki/[slug]": "Chi tiết Wiki", "/wiki/[slug]/edit": "Sửa Wiki", "/wiki/new": "Tạo Wiki", "/wishlist": "Danh sách yêu thích",
    }.get(route, "")) for route in evidence["web_routes"]]
    route_rows = []
    half = (len(route_items) + 1) // 2
    for i in range(half):
        left = route_items[i]
        right = route_items[i + half] if i + half < len(route_items) else ("", "")
        route_rows.append((left[0], left[1], right[0], right[1]))
    add_table(doc, ["Tuyến", "Vai trò", "Tuyến", "Vai trò"], route_rows,
              [1.25, 1.85, 1.25, 1.85], 9.4, raw_columns={0, 2})
    add_figure(doc, UI / "01-home.png", "4.1", "Trang chủ Spechub", "fig-4-1",
               "Ảnh giao diện trang chủ Spechub với khu vực giới thiệu, tìm kiếm và thẻ thiết bị.", page_break=True)
    add_body(doc, "Trang chủ ưu tiên hai hành động: bắt đầu tìm kiếm và đi vào các thiết bị nổi bật. Hero diễn đạt giá trị cốt lõi theo ngôn ngữ người dùng; thanh điều hướng giữ các đích Search, Devices, Compare, Recommend, AI và Wiki ở mức dễ thấy.")
    add_figure(doc, UI / "02-devices.png", "4.2", "Trang danh mục thiết bị", "fig-4-2",
               "Ảnh trang danh mục hiển thị tổng số model, hãng, loại, bộ lọc và danh sách thiết bị.", page_break=True)
    add_body(doc, "Trang danh mục kết hợp số liệu tổng quan, lọc và danh sách thẻ. Phân trang/lọc ở server giúp URL có thể chia sẻ; trạng thái rỗng và lỗi cần được phân biệt để người dùng biết nên đổi bộ lọc hay thử lại.")
    add_figure(doc, UI / "03-device-detail.png", "4.3", "Trang chi tiết thiết bị", "fig-4-3",
               "Ảnh trang chi tiết iPhone 16 Pro với tóm tắt phần cứng, biến thể và điểm đánh giá.", page_break=True)
    add_body(doc, "Trang chi tiết gom thông tin model, media, variant và phần cứng vào một ngữ cảnh. Việc hiển thị điểm cần đi kèm profile/trọng số hoặc diễn giải; nguồn nên đặt gần thuộc tính thay vì chỉ ở cuối trang để tăng khả năng kiểm chứng.")
    add_figure(doc, UI / "08-search.png", "4.4", "Kết quả tìm kiếm thiết bị và phần cứng", "fig-4-4",
               "Ảnh tìm kiếm từ khóa iPhone hiển thị thiết bị và kết quả phần cứng liên quan.", page_break=True)
    add_body(doc, "Kết quả tìm kiếm không giới hạn ở model mà có thể đưa ra thực thể phần cứng, giúp người nghiên cứu đi từ tên thương mại tới chipset hoặc linh kiện. Nhãn loại kết quả và slug rõ ràng giúp tránh nhầm lẫn giữa model, variant và component.")
    add_figure(doc, UI / "04-compare.png", "4.5", "Giao diện chọn thiết bị để so sánh", "fig-4-5",
               "Ảnh giao diện compare với vùng chọn thiết bị và trạng thái chờ đủ lựa chọn.", page_break=True)
    add_body(doc, "Giao diện so sánh bắt đầu bằng lựa chọn có kiểm soát thay vì hiển thị bảng trống. Sau khi đủ model, ma trận cần ghim cột tên, nhóm thuộc tính, làm nổi khác biệt và không điền giá trị suy đoán cho trường thiếu.")
    add_figure(doc, UI / "05-recommend.png", "4.6", "Quy trình khuyến nghị theo nhu cầu", "fig-4-6",
               "Ảnh quy trình khuyến nghị ba bước dựa trên nhu cầu và dữ liệu catalog.", page_break=True)
    add_body(doc, "Khuyến nghị dùng biểu mẫu ba bước để chuyển mục tiêu trừu tượng thành tiêu chí có cấu trúc trước khi gọi service. Cách này giảm độ mơ hồ, giúp giải thích vì sao một thiết bị được chọn và tạo fallback khi dịch vụ LLM không sẵn sàng.")
    add_figure(doc, UI / "06-ai.png", "4.7", "Giao diện trợ lý AI", "fig-4-7",
               "Ảnh giao diện trợ lý AI với vùng nhập câu hỏi và gợi ý chủ đề nghiên cứu.", page_break=True)
    add_body(doc, "Giao diện AI cần dành không gian cho citation, cảnh báo giới hạn và nút mở thực thể catalog. Nội dung hội thoại không được che khuất nguồn; câu trả lời nên hỗ trợ truy vấn tiếp theo nhưng không biến trạng thái chat thành nguồn sự thật cho lần truy xuất khác.")
    add_figure(doc, UI / "07-wiki.png", "4.8", "Giao diện Wiki công nghệ", "fig-4-8",
               "Ảnh giao diện Wiki công nghệ với vùng khám phá chủ đề và trạng thái dữ liệu tại thời điểm chụp.", page_break=True)
    add_body(doc, "Snapshot Wiki đang ở trạng thái chưa có bài xuất bản trong các bộ đếm hiển thị. Điều này được giữ nguyên như bằng chứng thực tế thay vì điền dữ liệu minh họa. Luồng tạo/sửa/duyệt đã có route và mô hình; độ phong phú nội dung là công việc dữ liệu tiếp theo.")
    add_figure(doc, UI / "09-login.png", "4.9", "Giao diện đăng nhập", "fig-4-9",
               "Ảnh biểu mẫu đăng nhập Spechub với trường email, mật khẩu và liên kết đăng ký.", page_break=True)
    add_body(doc, "Biểu mẫu đăng nhập tối giản, nhưng kiểm soát an toàn nằm ở API: validate, rate limit, kiểm tra băm, tạo session Redis và cookie. Thông báo sai thông tin cần đủ hữu ích nhưng không tiết lộ tài khoản có tồn tại.")

    add_heading(doc, "4.4. Kết quả API và dữ liệu", 2, "ch4-4")
    add_body(doc, "Kết quả phân tích tĩnh các bộ điều khiển xác định 186 điểm cuối API, bao gồm chức năng công khai, chức năng dành cho người dùng và chức năng quản trị. Số lượng này thể hiện phạm vi chức năng của API, nhưng không được sử dụng như tiêu chí duy nhất để đánh giá mức độ hoàn thiện. Trong mã nguồn, 45 điểm cuối được gắn vai trò ADMIN, 33 điểm cuối được gắn vai trò EDITOR và 3 điểm cuối được gắn vai trò MODERATOR; các điểm cuối còn lại sử dụng cơ chế xác thực chung, kiểm tra quyền sở hữu hoặc khai báo truy cập công khai.")
    add_table_caption(doc, "4.4", "Thống kê endpoint API", "tab-4-4")
    verbs = evidence["api"]["verbs"]
    roles = evidence["api"]["roles"]
    add_table(doc,
              ["Nhóm", "Số lượng", "Diễn giải"],
              [
                  ("Controller", evidence["api"]["controller_count"], "Nhóm endpoint theo miền"),
                  ("Endpoint tổng", evidence["api"]["endpoint_count"], "Route handler được phát hiện"),
                  ("GET", verbs["GET"], "Đọc catalog, health, dashboard và tra cứu"),
                  ("POST", verbs["POST"], "Tạo, đăng nhập, action, chat và đồng bộ"),
                  ("PATCH", verbs["PATCH"], "Cập nhật một phần và chuyển trạng thái"),
                  ("DELETE", verbs["DELETE"], "Xóa/thu hồi resource"),
                  ("PUT", verbs["PUT"], "Thay thế/cập nhật toàn phần ở trường hợp riêng"),
                  ("Role ADMIN", roles["ADMIN"], "Endpoint có decorator ADMIN"),
                  ("Role EDITOR", roles["EDITOR"], "Endpoint có decorator EDITOR"),
                  ("Role MODERATOR", roles["MODERATOR"], "Endpoint có decorator MODERATOR"),
              ], [1.5, 1.05, 3.65], 10.1)
    add_body(doc, "Lược đồ Prisma có 139 model và không khai báo enum Prisma; các trạng thái chủ yếu được biểu diễn bằng bảng tham chiếu hoặc trường chuỗi có kiểm soát ở tầng ứng dụng. Ưu điểm là dữ liệu tham chiếu có thể mở rộng; nhược điểm là cần bảo đảm validation và khóa ngoại để tránh giá trị trạng thái tùy ý.")

    add_heading(doc, "4.5. Kiểm thử và đánh giá", 2, "ch4-5")
    add_body(doc, "Việc đánh giá được thực hiện trên phiên bản mã nguồn ngày 05/08/2026. Toàn bộ kho mã nguồn được kiểm tra quy tắc viết mã, kiểu dữ liệu và khả năng biên dịch; bộ kiểm thử đơn vị của API được thực thi bằng Jest; gói chấm điểm được kiểm thử bằng Node test runner; lược đồ Prisma được kiểm tra bằng công cụ dòng lệnh Prisma. Mức sẵn sàng của API được ghi nhận khi PostgreSQL và Redis đang hoạt động trước khi đóng các cổng phát triển. Kết quả chỉ được sử dụng trong đúng phạm vi phép thử: khi chưa thực hiện kiểm thử tải thì chưa kết luận về hiệu năng trong môi trường vận hành; khi chưa có kiểm thử đầu cuối toàn tuyến thì kết quả kiểm thử đơn vị chưa đại diện cho toàn bộ luồng giao diện.")
    add_table_caption(doc, "4.5", "Bằng chứng kiểm thử và kiểm chứng", "tab-4-5")
    add_table(doc,
              ["Hạng mục", "Kết quả", "Phạm vi kết luận"],
              [
                  ("Kiểm tra quy tắc mã", "Đạt; 2 tác vụ lint hoàn tất", "Mã API và web không phát sinh lỗi từ cấu hình ESLint hiện hành"),
                  ("Kiểm tra kiểu dữ liệu", "Đạt; 7 tác vụ hoàn tất", "API, web, worker, AI service, database và gói chấm điểm được TypeScript chấp nhận"),
                  ("Biên dịch toàn dự án", "Đạt; 5 tác vụ build hoàn tất", "Next.js, NestJS, worker, AI service và gói chấm điểm biên dịch thành công"),
                  ("Kiểm thử API", "36 bộ đạt; 189 ca đạt; 0 thất bại", "Quy tắc, bộ điều khiển và dịch vụ có kiểm thử trong apps/api"),
                  ("Kiểm thử chấm điểm", "4 ca đạt; 0 thất bại", "Mốc tham chiếu, dữ liệu đo và tổng trọng số hồ sơ điểm"),
                  ("Prisma validate", "Lược đồ hợp lệ", "Cú pháp, quan hệ và cấu hình Prisma được CLI chấp nhận"),
                  ("Health tổng", "status = ok", "Tiến trình API phản hồi health"),
                  ("Readiness database", "ok", "API kết nối PostgreSQL trong môi trường kiểm tra"),
                  ("Readiness Redis", "ok", "API kết nối Redis trong môi trường kiểm tra"),
                  ("UI smoke", "9 màn hình chính được mở và chụp", "Route render với dữ liệu hiện có; không thay thế E2E assertion"),
                  ("Static inventory", "22 web route; 30 controller; 186 endpoint; 139 model", "Độ rộng mã nguồn tại commit/worktree đang phân tích"),
              ], [1.45, 2.15, 2.6], 9.8)
    add_listing(doc, "Tóm tắt kết quả đánh giá", [
        "Kiểm tra quy tắc mã, kiểu dữ liệu và biên dịch: đạt",
        "Kiểm thử API: 36 bộ đạt / 189 ca đạt / 0 ca thất bại",
        "Kiểm thử gói chấm điểm: 4 ca đạt / 0 ca thất bại",
        "Lược đồ Prisma: hợp lệ",
        "Trạng thái API: hoạt động bình thường",
        "Mức sẵn sàng: PostgreSQL đạt; Redis đạt",
        "Dữ liệu giao diện: 282 mẫu thiết bị / 39 hãng / 8 loại thiết bị",
    ])

    add_heading(doc, "4.6. Đánh giá mức độ đáp ứng", 2, "ch4-6")
    add_table_caption(doc, "4.6", "Đánh giá mức độ đáp ứng yêu cầu", "tab-4-6")
    add_table(doc,
              ["Nhóm yêu cầu", "Mức đáp ứng", "Kết quả / nhận xét"],
              [
                  ("Catalog và chi tiết", "Đạt ở mức hiện thực", "Route, API, schema và giao diện chi tiết; cần tiếp tục tăng độ phủ dữ liệu"),
                  ("Tìm kiếm và so sánh", "Đạt ở mức hiện thực", "Route/API/sequence và UI; cần benchmark relevance và tải"),
                  ("Khuyến nghị/AI", "Đạt kiến trúc và bề mặt", "Endpoint, UI và RAG design; cần bộ đánh giá chất lượng có chuẩn"),
                  ("Auth/RBAC", "Đạt ở mức hiện thực", "JWT, refresh session Redis, guard, role và unit test; cần pentest"),
                  ("Wiki/moderation", "Đạt mô hình và luồng", "Route/schema/controller; snapshot chưa có nội dung published đáng kể"),
                  ("Wishlist/alert", "Đạt mô hình và luồng", "Route, API, worker/notification design; cần E2E theo thời gian"),
                  ("B2B/billing", "Có nền tảng", "API key, usage, subscription, Stripe/webhook; cần thử nghiệm đối tác thật"),
                  ("Vận hành", "Đạt mức cơ bản", "Kiểm tra trạng thái, mức sẵn sàng, chỉ số giám sát và mã định danh yêu cầu; chưa có SLO và kết quả đo tải trong môi trường vận hành"),
              ], [1.45, 1.35, 3.4], 9.8)

    add_heading(doc, "4.7. Hạn chế", 2, "ch4-7")
    add_table_caption(doc, "4.7", "Hạn chế và tác động", "tab-4-7")
    add_table(doc,
              ["Hạn chế", "Tác động", "Ưu tiên xử lý"],
              [
                  ("Chưa có benchmark tải trong báo cáo", "Không thể công bố p95/p99 hoặc số người dùng đồng thời", "Cao: k6/Artillery theo search, compare, AI và auth"),
                  ("Độ phủ dữ liệu thay đổi theo seed/nguồn", "Một số category/Wiki có thể ít nội dung", "Cao: KPI completeness và pipeline nguồn"),
                  ("Chưa có E2E toàn tuyến được dẫn chứng", "Unit test không phát hiện mọi lỗi trình duyệt–API", "Cao: Playwright cho hành trình chính"),
                  ("AI chưa có bộ chấm chuẩn", "Khó định lượng độ đúng/citation", "Cao: golden set, RAG metrics và human review"),
                  ("Nhiều mô hình dữ liệu", "Migration và ownership phức tạp", "Trung bình: data dictionary, domain owner và migration policy"),
                  ("Dịch vụ ngoài có chế độ tùy chọn", "Hành vi khác nhau theo môi trường", "Trung bình: contract test và fallback matrix"),
                  ("Chưa chứng minh triển khai đa vùng", "Khả dụng/khôi phục thảm họa chưa được đánh giá", "Sau MVP: backup restore drill, RPO/RTO và chaos test"),
              ], [1.95, 2.1, 2.15], 9.7)
    add_body(doc, "Các hạn chế được trình bày nhằm phân biệt rõ giữa kết quả đã đạt được và mục tiêu phát triển tiếp theo. Mỗi hạn chế cần được chuyển thành phép đo, ca kiểm thử hoặc tiêu chí nghiệm thu cụ thể trong các giai đoạn tiếp theo.")

    add_heading(doc, "4.8. Kết luận chương", 2, "ch4-8")
    add_body(doc, "Chương 4 đã đối chiếu thiết kế với mã nguồn, giao diện và kết quả kiểm thử. Spechub đã xây dựng được phạm vi chức năng tương đối rộng, mô hình dữ liệu chi tiết và các thành phần vận hành cơ bản. Kiểm tra quy tắc mã, kiểu dữ liệu và biên dịch đều đạt; toàn bộ 189 ca kiểm thử API và 4 ca kiểm thử chấm điểm đều đạt; lược đồ Prisma hợp lệ; PostgreSQL và Redis đáp ứng kiểm tra mức sẵn sàng. Tuy nhiên, mức độ sẵn sàng triển khai thực tế còn phụ thuộc vào kết quả kiểm thử tải, kiểm thử đầu cuối, đánh giá chất lượng AI, chất lượng dữ liệu và cơ chế giám sát theo mục tiêu mức dịch vụ.")


def add_conclusion(doc):
    add_heading(doc, "KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", 1, "conclusion")
    add_heading(doc, "1. Kết luận", 2, "conclusion-1", toc=False)
    for text in [
        "Đồ án đã phân tích và tài liệu hóa một nền tảng nghiên cứu thiết bị có cấu trúc từ lớp trải nghiệm đến dữ liệu. Spechub giải quyết mối liên hệ giữa model, variant, phần cứng, nội dung, nguồn, người dùng và giá bằng 139 mô hình Prisma; cung cấp 22 route web và 186 endpoint API theo các miền catalog, search, AI, Wiki, alerts, commerce và vận hành.",
        "Kiến trúc sử dụng Next.js/React, NestJS/Fastify, Prisma/PostgreSQL, Redis và worker trong monorepo. Các quyết định quan trọng gồm tách published khỏi draft/raw, citation khỏi nội dung, notification khỏi delivery, session khỏi token và retrieval khỏi generation. Những ranh giới này giúp hệ thống có khả năng truy vết, kiểm thử và mở rộng tốt hơn.",
        "Kết quả đánh giá cho thấy kiểm tra quy tắc mã, kiểu dữ liệu và biên dịch đều đạt; 36 bộ với 189 ca kiểm thử API và 4 ca kiểm thử chấm điểm đều đạt; lược đồ Prisma hợp lệ; PostgreSQL và Redis đáp ứng kiểm tra mức sẵn sàng. Các giao diện chính có thể truy cập và hiển thị dữ liệu tại thời điểm đánh giá. Những kết quả này chưa thay thế cho kiểm thử tải, kiểm thử đầu cuối, kiểm thử xâm nhập hoặc đánh giá định lượng chất lượng AI, nhưng tạo được mốc tham chiếu cho các giai đoạn hoàn thiện tiếp theo.",
    ]:
        add_body(doc, text)
    add_heading(doc, "2. Hướng phát triển", 2, "conclusion-2", toc=False)
    add_numbered(doc, [
        "Xây dựng bộ KPI chất lượng catalog: độ đầy đủ theo category, freshness, conflict rate, citation coverage và thời gian review.",
        "Bổ sung Playwright E2E cho đăng nhập, tìm kiếm, so sánh, khuyến nghị, Wiki review và cảnh báo giá; chạy trong CI với dữ liệu seed ổn định.",
        "Thiết lập kiểm thử tải và SLO cho search, device detail, compare, auth refresh và AI; theo dõi p50/p95/p99, error rate và saturation.",
        "Tạo golden set cho AI/RAG, đo context recall, citation precision, faithfulness và chất lượng khuyến nghị; có đánh giá con người định kỳ.",
        "Hoàn thiện pipeline ingestion theo từng nguồn với parser contract test, cảnh báo thay đổi DOM, ưu tiên nguồn chính thức và chính sách tuân thủ.",
        "Nâng cấp an toàn bằng threat-model review, dependency scanning, secret scanning, pentest, backup/restore drill và diễn tập thu hồi khóa; đồng thời mở rộng API B2B và Catalog Studio từ phản hồi của đối tác và biên tập viên thực tế.",
    ], compact=True, font_size=10.5)


def add_references(doc):
    add_heading(doc, "TÀI LIỆU THAM KHẢO", 1, "references")
    refs = [
        "[1] Kho mã nguồn Spechub, tệp README và tài liệu kiến trúc nội bộ của dự án, truy cập ngày 05/08/2026.",
        "[2] Next.js Documentation, https://nextjs.org/docs, truy cập ngày 05/08/2026.",
        "[3] React Documentation, https://react.dev/, truy cập ngày 05/08/2026.",
        "[4] NestJS Documentation, https://docs.nestjs.com/, truy cập ngày 05/08/2026.",
        "[5] Fastify Documentation, https://fastify.dev/docs/latest/, truy cập ngày 05/08/2026.",
        "[6] Prisma ORM Documentation, https://www.prisma.io/docs/orm, truy cập ngày 05/08/2026.",
        "[7] PostgreSQL 16 Documentation, https://www.postgresql.org/docs/16/, truy cập ngày 05/08/2026.",
        "[8] pgvector Documentation, https://github.com/pgvector/pgvector, truy cập ngày 05/08/2026.",
        "[9] Redis Documentation, https://redis.io/docs/latest/, truy cập ngày 05/08/2026.",
        "[10] Meilisearch Documentation, https://www.meilisearch.com/docs, truy cập ngày 05/08/2026.",
        "[11] OWASP Application Security Verification Standard, https://owasp.org/www-project-application-security-verification-standard/, truy cập ngày 05/08/2026.",
        "[12] OWASP API Security Top 10, https://owasp.org/www-project-api-security/, truy cập ngày 05/08/2026.",
        "[13] OpenAI, Retrieval and model prompting documentation, https://platform.openai.com/docs/, truy cập ngày 05/08/2026.",
        "[14] Stripe Documentation – Webhooks and subscriptions, https://docs.stripe.com/, truy cập ngày 05/08/2026.",
        "[15] W3C, Web Content Accessibility Guidelines (WCAG) 2.2, https://www.w3.org/TR/WCAG22/, truy cập ngày 05/08/2026.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(ref)
        set_run_font(r, 10.5)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)


def model_group(name):
    n = name.lower()
    if any(k in n for k in ("user", "wishlist", "alert", "notification", "subscription", "api_key", "api_usage", "affiliate", "audit", "webhook", "price")):
        return "Người dùng / thương mại"
    if any(k in n for k in ("wiki", "source", "citation", "raw_page", "embedding", "search_log", "ai_", "comment")):
        return "Nội dung / AI / nguồn"
    if any(k in n for k in ("chip", "cpu", "gpu", "npu", "display", "camera", "battery", "memory", "storage", "benchmark", "score", "operating")):
        return "Phần cứng / điểm"
    return "Catalog / tham chiếu"


def add_appendices(doc, evidence):
    add_heading(doc, "PHỤ LỤC A. DANH MỤC API", 1, "appendix-a")
    add_body(doc, "Bảng A.1 được tổng hợp từ các bộ điều khiển trong mã nguồn. Trường hợp tiền tố rỗng hoặc bộ điều khiển không có hàm xử lý trực tiếp vẫn được giữ lại để phản ánh đúng cấu trúc mô-đun. Thông tin chi tiết về từng điểm cuối API có thể được tra cứu bằng Swagger trong môi trường phát triển.")
    add_table_caption(doc, "A.1", "Danh mục controller và endpoint", "tab-a-1")
    rows = []
    for idx, ctrl in enumerate(evidence["api"]["controllers"], 1):
        name = Path(ctrl["file"]).stem.replace(".controller", "")
        prefix = ctrl.get("prefix") or "(gốc)"
        methods = {}
        for ep in ctrl.get("endpoints", []):
            methods[ep["verb"]] = methods.get(ep["verb"], 0) + 1
        method_text = ", ".join(f"{k}:{v}" for k, v in sorted(methods.items())) or "–"
        rows.append((idx, name, prefix, ctrl["endpoint_count"], method_text))
    add_table(doc, ["STT", "Bộ điều khiển", "Tiền tố", "Số điểm cuối API", "Phân bố phương thức"], rows,
              [0.45, 1.75, 1.85, 0.8, 1.35], 9.0, raw_columns={1, 2, 4})
    add_body(doc, "Các nhóm endpoint có bề mặt lớn gồm affiliate/price, AI, catalog, Wiki/moderation, notification, admin và billing. Khi mở rộng, cần ưu tiên tính nhất quán của DTO, status code, pagination và quyền hơn là tăng số route.")

    add_heading(doc, "PHỤ LỤC B. DANH MỤC MÔ HÌNH DỮ LIỆU", 1, "appendix-b")
    schema = SCHEMA_PATH.read_text(encoding="utf-8")
    models = re.findall(r"(?m)^model\s+([A-Za-z0-9_]+)\s*\{", schema)
    add_body(doc, f"Lược đồ vật lý có {len(models)} model Prisma. Bảng B.1 liệt kê toàn bộ tên model theo thứ tự xuất hiện trong schema và nhóm phân tích gần đúng. Nhóm chỉ phục vụ đọc báo cáo; quan hệ thật được định nghĩa trong Prisma.")
    add_table_caption(doc, "B.1", "Danh sách mô hình Prisma", "tab-b-1")
    model_rows = [(i + 1, name, model_group(name)) for i, name in enumerate(models)]
    add_table(doc, ["STT", "Tên mô hình", "Nhóm phân tích"], model_rows,
              [0.65, 3.0, 2.55], 8.8, raw_columns={1})

    add_heading(doc, "PHỤ LỤC C. HƯỚNG DẪN VẬN HÀNH VÀ KIỂM THỬ", 1, "appendix-c")
    add_body(doc, "Các lệnh dưới đây mang tính quy trình; tên script cụ thể phải đối chiếu package.json tại thời điểm chạy. Không đưa bí mật hoặc giá trị production vào tài liệu. Tệp .env chỉ được tạo từ mẫu và quản lý ngoài kiểm soát phiên bản.")
    add_table_caption(doc, "C.1", "Ma trận lệnh vận hành và mục đích", "tab-c-1")
    add_table(doc,
              ["Bước", "Lệnh tham chiếu", "Mục đích / kết quả mong đợi"],
              [
                  ("Cài đặt", "pnpm install --frozen-lockfile", "Cài đúng lockfile trong workspace"),
                  ("Hạ tầng", "pnpm docker:up hoặc docker compose up", "Khởi động PostgreSQL, Redis và dịch vụ tùy chọn"),
                  ("Prisma", "pnpm --filter database prisma validate", "Lược đồ hợp lệ trước generate/migrate"),
                  ("Generate", "pnpm --filter database prisma generate", "Tạo Prisma client đúng schema"),
                  ("Migration", "Prisma migrate theo môi trường", "Áp migration có review; không dùng reset trên dữ liệu cần giữ"),
                  ("Seed", "Script seed của packages/database", "Tạo dữ liệu phát triển có thể tái lập"),
                  ("Dev", "pnpm dev", "Chạy pipeline web/API/worker theo cấu hình"),
                  ("Test API", "pnpm --filter api test", "Chạy Jest và báo suite/test thất bại"),
                  ("Typecheck", "pnpm typecheck", "Kiểm tra kiểu toàn workspace"),
                  ("Lint", "pnpm lint", "Kiểm tra quy ước và lỗi tĩnh"),
                  ("Health", "GET /api/v1/health/ready", "DB/Redis phải báo ready trước nhận lưu lượng"),
                  ("Backup", "Quy trình pg_dump/restore đã phê duyệt", "Kiểm tra khả năng phục hồi, không chỉ tạo bản sao"),
              ], [0.85, 2.25, 3.1], 9.5, raw_columns={1})
    add_heading(doc, "C.1. Checklist trước khi triển khai", 2, "appendix-c-1", toc=False)
    add_bullets(doc, [
        "Tất cả migration đã review, backup có thể phục hồi và biến môi trường bắt buộc đã được kiểm tra.",
        "Test, typecheck, lint và build đều đạt trên cùng commit; image/container có version bất biến.",
        "CORS, cookie, JWT secret, webhook secret, API key hashing và rate limit đúng môi trường.",
        "Health/readiness/metrics, log redaction, request ID, dashboard và cảnh báo đã hoạt động.",
        "Chỉ mục tìm kiếm và véc-tơ nhúng được quản lý phiên bản; tiến trình nền bảo đảm tính lũy đẳng; cơ chế thử lại và hàng đợi lỗi được giám sát.",
        "Có kế hoạch rollback, người chịu trách nhiệm và tiêu chí dừng triển khai.",
    ])

    add_heading(doc, "PHỤ LỤC D. HƯỚNG DẪN SỬ DỤNG TÓM TẮT", 1, "appendix-d")
    add_body(doc, "Hướng dẫn này mô tả hành trình ở mức người dùng và không yêu cầu hiểu cấu trúc kỹ thuật.")
    add_table_caption(doc, "D.1", "Tác vụ người dùng thường gặp", "tab-d-1")
    add_table(doc,
              ["Tác vụ", "Cách thực hiện", "Lưu ý"],
              [
                  ("Tìm thiết bị", "Mở Search/Devices, nhập từ khóa, chọn loại/hãng và bộ lọc", "Bỏ bớt filter nếu không có kết quả; kiểm tra alias"),
                  ("Xem chi tiết", "Chọn thẻ model để xem variant, phần cứng, điểm và nguồn", "Phân biệt model với variant; trường thiếu không phải giá trị 0"),
                  ("So sánh", "Thêm từ hai thiết bị vào Compare", "Ưu tiên cùng category; mở nguồn khi giá trị mâu thuẫn"),
                  ("Nhận khuyến nghị", "Hoàn thành các bước nhu cầu/ngân sách/ưu tiên", "Kết quả là hỗ trợ quyết định, không thay thế kiểm tra thực tế"),
                  ("Hỏi AI", "Đặt câu hỏi cụ thể và xem citation", "Không coi câu không có nguồn là dữ kiện đã xác minh"),
                  ("Wishlist/cảnh báo", "Đăng nhập, lưu variant và đặt target price", "Kiểm tra tiền tệ, kênh nhận và tùy chọn thông báo"),
                  ("Đóng góp Wiki", "Tạo/sửa revision, thêm nguồn và gửi duyệt", "Bản sửa không xuất bản ngay; có thể được yêu cầu bổ sung"),
                  ("Dùng API", "Tạo API key theo gói, lưu khóa an toàn và gọi endpoint trong scope", "Khóa chỉ hiển thị rõ một lần; theo dõi quota/usage"),
              ], [1.35, 3.05, 1.8], 9.8)
    add_heading(doc, "D.1. Xử lý sự cố cơ bản", 2, "appendix-d-1", toc=False)
    add_bullets(doc, [
        "Nếu trang báo offline: kiểm tra kết nối, thử tải lại; dữ liệu cache có thể không phải giá mới nhất.",
        "Nếu đăng nhập hết hạn: thực hiện đăng nhập lại; không gửi token hoặc ảnh chụp token cho người khác.",
        "Nếu kết quả sai: mở trích dẫn nguồn, ghi URL, thuộc tính và biến thể liên quan, sau đó gửi phản hồi kèm nguồn kiểm chứng.",
        "Nếu cảnh báo không đến: kiểm tra alert active, ngưỡng, kênh nhận, quiet hours và trạng thái notification.",
        "Nếu API trả 429: đọc header quota/retry, giảm tần suất và không tạo thêm key để né hạn mức.",
    ])


def add_document_settings(doc):
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "false")
    compat = settings.find(qn("w:compat"))
    if compat is None:
        compat = OxmlElement("w:compat")
        settings.append(compat)
    balance = OxmlElement("w:doNotUseHTMLParagraphAutoSpacing")
    compat.append(balance)


def document_from_reference_template():
    """Tạo bản làm việc từ mẫu chính và chỉ thay phần thân tài liệu.

    Việc giữ lại sectPr cuối, theme, styles, numbering và các package part của
    mẫu bảo đảm báo cáo mới vẫn có nguồn gốc trực tiếp từ tài liệu tham chiếu.
    Các đoạn, bảng và hình thuộc đề tài mẫu được loại khỏi body trước khi thêm
    nội dung Spechub.
    """
    doc = Document(TEMPLATE)
    body = doc._element.body
    final_sect_pr = body.sectPr
    for child in list(body):
        if child is not final_sect_pr:
            body.remove(child)
    return doc


def prune_reference_artifacts(path: Path):
    """Loại các part của đề tài mẫu không còn được nội dung mới tham chiếu."""
    rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    office_rel_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    word_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    content_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    document_rels_path = "word/_rels/document.xml.rels"
    content_types_path = "[Content_Types].xml"
    tmp = path.with_name(path.stem + ".cleaning.docx")

    with ZipFile(path, "r") as source:
        payloads = {name: source.read(name) for name in source.namelist()}

    document = etree.fromstring(payloads["word/document.xml"])
    referenced_ids = set(
        document.xpath(
            "//@r:id | //@r:embed | //@r:link",
            namespaces={"r": office_rel_ns},
        )
    )
    has_footnotes = bool(
        document.xpath(".//w:footnoteReference", namespaces={"w": word_ns})
    )
    has_endnotes = bool(
        document.xpath(".//w:endnoteReference", namespaces={"w": word_ns})
    )
    has_comments = bool(
        document.xpath(".//w:commentReference", namespaces={"w": word_ns})
    )

    rels = etree.fromstring(payloads[document_rels_path])
    deleted_parts: set[str] = set()
    prunable = {"header", "footer", "image", "hyperlink"}
    optional = {
        "footnotes": has_footnotes,
        "endnotes": has_endnotes,
        "comments": has_comments,
        "commentsExtended": has_comments,
        "commentsIds": has_comments,
        "people": has_comments,
    }
    for rel in list(rels):
        rel_id = rel.get("Id", "")
        rel_type = rel.get("Type", "").rsplit("/", 1)[-1]
        target = rel.get("Target", "")
        remove = (rel_type in prunable and rel_id not in referenced_ids) or (
            rel_type in optional and not optional[rel_type]
        ) or rel_type == "customXml"
        if not remove:
            continue
        rels.remove(rel)
        if target and not target.startswith(("http://", "https://")):
            deleted_parts.add(posixpath.normpath(posixpath.join("word", target)))
    payloads[document_rels_path] = etree.tostring(
        rels, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    # Part phụ thuộc của header/footer cũ và các cấu trúc mẫu không còn dùng.
    for part in list(deleted_parts):
        directory, filename = posixpath.split(part)
        deleted_parts.add(posixpath.join(directory, "_rels", filename + ".rels"))
    if not has_footnotes:
        deleted_parts.add("word/footnotes.xml")
    if not has_endnotes:
        deleted_parts.add("word/endnotes.xml")
    if not has_comments:
        deleted_parts.update(
            name
            for name in payloads
            if name.startswith(("word/comments", "word/people"))
        )
    deleted_parts.update(name for name in payloads if name.startswith("customXml/"))

    # Chỉ giữ media còn được một relationship của package sau làm sạch tham chiếu.
    referenced_parts: set[str] = set()
    for rel_path, data in payloads.items():
        if not rel_path.endswith(".rels") or rel_path in deleted_parts:
            continue
        if rel_path == document_rels_path:
            rel_root = rels
        else:
            try:
                rel_root = etree.fromstring(data)
            except etree.XMLSyntaxError:
                continue
        if "/_rels/" in rel_path:
            source_dir = rel_path.split("/_rels/", 1)[0]
        else:
            source_dir = ""
        for rel in rel_root.findall(f"{{{rel_ns}}}Relationship"):
            target = rel.get("Target", "")
            if not target or target.startswith(("http://", "https://")):
                continue
            referenced_parts.add(posixpath.normpath(posixpath.join(source_dir, target)))
    deleted_parts.update(
        name
        for name in payloads
        if name.startswith("word/media/") and name not in referenced_parts
    )

    content_types = etree.fromstring(payloads[content_types_path])
    for node in list(content_types):
        if etree.QName(node).namespace != content_ns:
            continue
        part_name = node.get("PartName", "").lstrip("/")
        if part_name in deleted_parts or part_name.startswith("customXml/"):
            content_types.remove(node)
    payloads[content_types_path] = etree.tostring(
        content_types, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    with ZipFile(tmp, "w", ZIP_DEFLATED) as target:
        for name, data in payloads.items():
            if name not in deleted_parts:
                target.writestr(name, data)
    tmp.replace(path)


def build(output_path):
    evidence = json.loads(EVIDENCE_PATH.read_text(encoding="utf-8"))
    doc = document_from_reference_template()
    configure_styles(doc)
    set_core_properties(doc)
    add_document_settings(doc)

    configure_section(doc.sections[0], numbered=False)
    set_page_border(doc.sections[0])
    cover(doc, inner=False)

    inner = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(inner, numbered=False)
    cover(doc, inner=True)

    prelim = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(prelim, numbered=True, fmt="lowerRoman", start=1)
    add_summary(doc)
    add_declaration(doc)
    add_acknowledgements(doc)
    add_toc(doc)
    add_front_lists(doc)

    body = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(body, numbered=True, fmt="decimal", start=1)
    add_opening(doc)
    add_chapter_1(doc)
    add_chapter_2(doc, evidence)
    add_chapter_3(doc)
    add_chapter_4(doc, evidence)
    add_conclusion(doc)
    add_references(doc)
    add_appendices(doc, evidence)

    doc.save(output_path)
    prune_reference_artifacts(Path(output_path))


def main():
    global PAGE_MAP, MARKERS
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", required=True)
    parser.add_argument("--page-map")
    parser.add_argument("--markers", action="store_true")
    args = parser.parse_args()
    PAGE_MAP = load_page_map(args.page_map)
    MARKERS = args.markers
    build(Path(args.output))
    print(args.output)


if __name__ == "__main__":
    main()
